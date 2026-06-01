# -*- coding: utf-8 -*-
"""
Génère le guide utilisateur « Meeting Assistant » au format Word (.docx),
en français, pour des utilisateurs NON TECHNIQUES (aucun jargon).

Pour chaque étape d'installation et chaque fonctionnalité, une phrase EN ROUGE
décrit la capture d'écran à insérer (quoi montrer, quoi entourer).

Le script installe python-docx si nécessaire, puis écrit :
    docs/Guide_Utilisateur_Meeting_Assistant.docx

Usage :
    python docs/generate_user_guide_docx.py
"""

import os
import subprocess
import sys

# --------------------------------------------------------------------------- #
#  Dépendance : python-docx (auto-installation si absente)
# --------------------------------------------------------------------------- #
try:
    import docx  # noqa: F401
except ImportError:
    print("python-docx absent — installation…")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------- #
#  Dossier des captures d'écran
# --------------------------------------------------------------------------- #
SHOTS_DIR = (
    r"C:\Users\MelissaBELGUITAR\OneDrive - YELE CONSULTING"
    "\\Images\\Captures d’écran"
)


def shot_path(filename):
    p = os.path.join(SHOTS_DIR, filename)
    return p if os.path.exists(p) else None

# --------------------------------------------------------------------------- #
#  Charte graphique
# --------------------------------------------------------------------------- #
BRAND      = RGBColor(0x1E, 0x40, 0xAF)   # bleu profond (titres)
BRAND_LT   = RGBColor(0x3B, 0x82, 0xF6)   # bleu clair (kicker / puces)
DARK       = RGBColor(0x0F, 0x17, 0x2A)   # texte principal
MUTED      = RGBColor(0x55, 0x65, 0x77)   # texte secondaire
RED        = RGBColor(0xC0, 0x10, 0x10)   # phrase capture (rouge)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
AMBER_TXT  = RGBColor(0x92, 0x5A, 0x05)   # texte des encadrés conseils

FONT = "Calibri"


# --------------------------------------------------------------------------- #
#  Helpers bas niveau
# --------------------------------------------------------------------------- #
def _set_cell_bg(cell, hex_color):
    """Colore le fond d'une cellule de tableau (manip XML)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _shade_paragraph(p, hex_color):
    """Applique un fond de couleur derrière un paragraphe (encadré plein)."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _border_paragraph(p, hex_color="C7D2FE", size=8, space=6):
    """Bordure complète autour d'un paragraphe (cadre)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), hex_color)
        pbdr.append(el)
    pPr.append(pbdr)


def run(p, text, size=11, color=DARK, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def para(doc, space_before=0, space_after=6, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    return p


# --------------------------------------------------------------------------- #
#  Composants de contenu
# --------------------------------------------------------------------------- #
def h1(doc, number, text):
    """Titre de partie : kicker numéroté + libellé souligné."""
    p = para(doc, space_before=18, space_after=4)
    run(p, f"PARTIE {number}", 12, BRAND_LT, bold=True)
    p2 = para(doc, space_before=0, space_after=10)
    run(p2, text, 22, BRAND, bold=True)
    pPr = p2._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E40AF")
    pbdr.append(bottom)
    pPr.append(pbdr)


def h2(doc, kicker, title):
    """Sous-titre d'étape / fonctionnalité."""
    if kicker:
        p = para(doc, space_before=14, space_after=1)
        run(p, kicker.upper(), 10, BRAND_LT, bold=True)
        p2 = para(doc, space_before=0, space_after=5)
    else:
        p2 = para(doc, space_before=14, space_after=5)
    run(p2, title, 15, DARK, bold=True)


def bullet(doc, segments, level=0):
    """Puce. `segments` = texte simple OU liste de fragments (texte, gras)
    pour mettre en gras les libellés de boutons / champs au fil de la phrase."""
    p = para(doc, space_after=4)
    p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    glyph = "•  " if level == 0 else "–  "
    run(p, glyph, 11, BRAND_LT if level == 0 else MUTED, bold=True)
    txt_color = DARK if level == 0 else MUTED
    if isinstance(segments, str):
        segments = [(segments, False)]
    for seg in segments:
        text, bold = seg
        run(p, text, 11, txt_color, bold=bold)
    return p


def shot(doc, text, image=None, width=5.5):
    """Insère une capture d'écran.

    - Si `image` est un chemin de fichier existant : insère l'image centrée
      + une légende en italique gris.
    - Sinon : phrase EN ROUGE « Capture à insérer — … » (placeholder pour
      les emplacements où aucune capture n'est encore disponible).
    """
    img_path = shot_path(image) if image else None
    if img_path:
        # Image centrée
        pp = para(doc, space_before=6, space_after=2,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        r = pp.add_run()
        r.add_picture(img_path, width=Inches(width))
        # Légende grise centrée
        cap = para(doc, space_before=0, space_after=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        run(cap, "Figure — ", 9, MUTED, bold=True, italic=True)
        run(cap, text, 9, MUTED, italic=True)
        return pp
    # Placeholder rouge (pas de capture disponible)
    p = para(doc, space_before=4, space_after=10)
    p.paragraph_format.left_indent = Inches(0.3)
    run(p, "📷  Capture à insérer — ", 11, RED, bold=True)
    run(p, text, 11, RED)
    return p


def tip(doc, text):
    """Encadré conseil (fond ambre clair)."""
    p = para(doc, space_before=2, space_after=10)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    _shade_paragraph(p, "FEF3C7")
    _border_paragraph(p, "F59E0B", size=6, space=6)
    run(p, "💡  Conseil — ", 11, AMBER_TXT, bold=True)
    run(p, text, 11, AMBER_TXT)
    return p


def body(doc, text, italic=False, color=DARK):
    p = para(doc, space_after=6)
    run(p, text, 11, color, italic=italic)
    return p


def faq_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(4.1)
    hdr = table.rows[0].cells
    for i, t in enumerate(("Situation", "Que faire")):
        _set_cell_bg(hdr[i], "1E40AF")
        c = hdr[i].paragraphs[0]
        c.paragraph_format.space_after = Pt(2)
        c.paragraph_format.space_before = Pt(2)
        run(c, t, 11, WHITE, bold=True)
    for idx, (a, b) in enumerate(rows):
        cells = table.add_row().cells
        if idx % 2 == 0:
            _set_cell_bg(cells[0], "F1F5F9")
            _set_cell_bg(cells[1], "F1F5F9")
        pa = cells[0].paragraphs[0]
        pa.paragraph_format.space_after = Pt(2)
        pa.paragraph_format.space_before = Pt(2)
        run(pa, a, 10.5, DARK, bold=True)
        pb = cells[1].paragraphs[0]
        pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.space_before = Pt(2)
        run(pb, b, 10.5, DARK)


def page_break(doc):
    doc.add_page_break()


# --------------------------------------------------------------------------- #
#  Construction du document
# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    # ----------------------------------------------------------------- #
    #  Page de garde
    # ----------------------------------------------------------------- #
    for _ in range(3):
        para(doc, space_after=0)
    p = para(doc, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "GUIDE D'INSTALLATION ET D'UTILISATION", 14, BRAND_LT, bold=True)
    p = para(doc, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Meeting Assistant", 40, BRAND, bold=True)
    p = para(doc, space_after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Enregistrez vos réunions · Transcription automatique · "
           "Compte rendu Word généré pour vous", 14, MUTED)
    p = para(doc, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Aucune compétence technique requise", 12, MUTED, italic=True)
    p = para(doc, space_before=40, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Yele Consulting", 12, DARK, bold=True)

    p = para(doc, space_before=36, space_after=6)
    _shade_paragraph(p, "EFF4FB")
    _border_paragraph(p, "607DA8", size=6, space=8)
    run(p, "Comment lire ce guide  ·  ", 11, BRAND, bold=True)
    run(p, "Les phrases ", 11, DARK)
    run(p, "en rouge", 11, RED, bold=True)
    run(p, " indiquent la capture d'écran à insérer à cet endroit : elles "
           "décrivent ce qu'il faut montrer et ce qu'il faut entourer. "
           "Remplacez chaque phrase rouge par votre capture une fois le "
           "guide complété.", 11, DARK)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 1 — Découvrir l'application
    # ----------------------------------------------------------------- #
    h1(doc, 1, "Découvrir l'application")

    h2(doc, "À quoi ça sert", "Un assistant qui rédige vos comptes rendus")
    body(doc, "Meeting Assistant transforme une réunion en compte rendu "
              "Word, automatiquement :")
    bullet(doc, "Vous enregistrez la réunion en direct, ou vous importez un "
                "fichier audio (ou un transcript déjà produit).")
    bullet(doc, "L'application écrit toute seule ce qui a été dit "
                "(la transcription).")
    bullet(doc, "Elle reconnaît les différentes personnes qui parlent "
                "(« qui a dit quoi »).")
    bullet(doc, "Elle rédige un compte rendu Word clair : résumé, décisions, "
                "actions à mener.")
    bullet(doc, "Tout se passe sur votre ordinateur — vos données ne partent "
                "pas sur Internet.", level=1)
    bullet(doc, "En option, elle peut afficher vos réunions Outlook.", level=1)
    tip(doc, "Gain de temps : plus besoin de prendre des notes pendant la "
             "réunion ni de rédiger le compte rendu à la main.")
    shot(doc, "exemple de compte rendu généré automatiquement, ouvert dans "
              "l'application — voilà ce que vous allez obtenir.",
         image="Capture d'écran 2026-05-22 140103.png", width=6.0)

    h2(doc, "Avant de commencer", "Ce dont vous avez besoin")
    bullet(doc, "Un ordinateur sous Windows 10 ou Windows 11.")
    bullet(doc, "Environ 3 à 4 Go d'espace libre sur le disque.")
    bullet(doc, "Un microphone (celui de l'ordinateur ou un casque) pour "
                "enregistrer en direct.")
    bullet(doc, "Une connexion Internet, uniquement la première fois que "
                "vous ouvrez l'application.")
    bullet(doc, "Ensuite, l'application fonctionne sans Internet.", level=1)
    bullet(doc, "Aucun compte ni mot de passe n'est nécessaire pour démarrer.")
    tip(doc, "La toute première ouverture télécharge environ 2,3 Go : prévoyez "
             "une bonne connexion et un peu de patience. Cela n'arrive qu'une "
             "seule fois.")

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 2 — Installer l'application
    # ----------------------------------------------------------------- #
    h1(doc, 2, "Installer l'application")
    body(doc, "Six étapes simples, du fichier d'installation au premier "
              "démarrage.", italic=True, color=MUTED)

    h2(doc, "Étape 1 sur 6", "Récupérer le fichier d'installation")
    bullet(doc, [("Vous recevez un fichier d'installation nommé ", False),
                 ("MeetingAssistant-Setup.exe", True)])
    bullet(doc, "Enregistrez-le dans un dossier facile à retrouver "
                "(par exemple « Téléchargements »).")
    bullet(doc, "Le fichier est volumineux : son téléchargement peut prendre "
                "quelques minutes.")
    shot(doc, "le fichier « MeetingAssistant-Setup.exe » reçu — ici dans "
              "OneDrive, mais l'idée est la même dans Téléchargements ou en "
              "pièce jointe d'un e-mail.",
         image="Capture d'écran 2026-05-21 172502.png", width=3.2)

    h2(doc, "Étape 2 sur 6", "Lancer l'installation (message de Windows)")
    bullet(doc, "Double-cliquez sur le fichier téléchargé.")
    bullet(doc, "Windows peut afficher un écran bleu « Windows a protégé "
                "votre ordinateur ».")
    bullet(doc, [("Cliquez sur ", False),
                 ("Informations complémentaires", True), (".", False)])
    bullet(doc, [("Puis sur ", False),
                 ("Exécuter quand même", True), (".", False)])
    bullet(doc, "C'est normal pour une application récente : le fichier "
                "n'est pas dangereux.", level=1)
    shot(doc, "Windows demande de confirmer qu'on fait confiance au "
              "fichier — cliquez sur « Conserver quand même » (ou "
              "« Exécuter quand même » selon votre version de Windows).",
         image="Capture d'écran 2026-05-19 110330.png", width=5.0)

    h2(doc, "Étape 3 sur 6", "Suivre l'assistant d'installation")
    bullet(doc, "Une fenêtre d'installation s'ouvre.")
    bullet(doc, "Vous pouvez tout laisser tel quel.")
    bullet(doc, "Aucun mot de passe administrateur n'est demandé.")
    bullet(doc, [("Cliquez sur ", False), ("Installer", True),
                 (" et patientez quelques instants.", False)])
    shot(doc, "l'assistant d'installation : laissez le dossier proposé "
              "et cliquez sur « Installer ».",
         image="Capture d'écran 2026-05-19 110918.png", width=4.8)

    h2(doc, "Étape 4 sur 6", "Terminer l'installation")
    bullet(doc, "Un message confirme que l'installation est terminée.")
    bullet(doc, "Un raccourci « Meeting Assistant » apparaît sur le Bureau "
                "et dans le menu Démarrer.")
    bullet(doc, [("Laissez la case ", False),
                 ("« Lancer Meeting Assistant »", True),
                 (" cochée, puis cliquez sur ", False),
                 ("Fermer", True), (".", False)])
    shot(doc, "le dernier écran de l'installation — cliquez sur « Fermer » "
              "(la case « Lancer Meeting Assistant » ouvre l'application "
              "immédiatement).",
         image="Capture d'écran 2026-05-21 181337.png", width=4.5)

    h2(doc, "Étape 5 sur 6",
       "Première ouverture : préparation de l'application")
    bullet(doc, [("Double-cliquez sur l'icône ", False),
                 ("Meeting Assistant", True), (".", False)])
    bullet(doc, "Une fenêtre « Préparation de Meeting Assistant » s'affiche.")
    bullet(doc, "L'application télécharge ses outils (environ 2,3 Go) avec "
                "une barre de progression.")
    bullet(doc, "Cela dure de 5 à 30 minutes selon votre connexion.")
    bullet(doc, "Ne fermez pas la fenêtre — cela n'arrive qu'une seule fois.",
           level=1)
    tip(doc, "Si le téléchargement est interrompu (coupure Internet), il "
             "reprend tout seul là où il s'était arrêté au prochain démarrage.")
    shot(doc, "la fenêtre « Préparation de Meeting Assistant » — barre de "
              "progression, numéro de fichier en cours et pourcentage.",
         image="Capture d'écran 2026-05-22 134306.png", width=5.0)

    h2(doc, "Étape 6 sur 6", "L'application s'ouvre")
    bullet(doc, "Une fois la préparation finie, un écran « Meeting Assistant "
                "— Démarrage… » s'affiche brièvement.")
    bullet(doc, "L'application s'ouvre ensuite toute seule.")
    bullet(doc, "Les fois suivantes, l'ouverture sera beaucoup plus rapide.")
    shot(doc, "le petit écran de démarrage « Meeting Assistant — "
              "Démarrage… » qui apparaît brièvement à chaque ouverture.",
         image="Capture d'écran 2026-05-21 181645.png", width=4.0)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 3 — Découvrir l'écran principal
    # ----------------------------------------------------------------- #
    h1(doc, 3, "Découvrir l'écran principal")

    h2(doc, "Vue d'ensemble", "Se repérer dans l'application")
    bullet(doc, "À gauche, la barre latérale donne accès à tout :")
    bullet(doc, [("Nouvelle réunion", True),
                 (" — pour démarrer une réunion.", False)], level=1)
    bullet(doc, [("Comptes rendus", True),
                 (" — tous vos comptes rendus déjà produits.", False)],
           level=1)
    bullet(doc, [("Agenda", True),
                 (" — vos réunions Outlook, si l'agenda est connecté.",
                  False)], level=1)
    bullet(doc, [("Dossiers", True),
                 (" — pour classer vos réunions par client ou projet.",
                  False)], level=1)
    bullet(doc, [("En bas à gauche : ", False), ("Paramètres", True),
                 (", « Contactez-nous » et le bouton clair / sombre.",
                  False)])
    bullet(doc, "Au centre : votre agenda du jour et vos prochaines réunions.")
    shot(doc, "l'écran d'accueil. La barre latérale à gauche regroupe tous "
              "les accès ; au centre, l'agenda du jour et les prochaines "
              "réunions.",
         image="Capture d'écran 2026-05-22 104919.png", width=6.0)

    h2(doc, "Réglages", "Ouvrir les Paramètres")
    bullet(doc, [("Cliquez sur ", False), ("Paramètres", True),
                 (" en bas de la barre latérale.", False)])
    bullet(doc, "Vous y trouvez : la clé pour le compte rendu rapide en "
                "ligne, les options d'arrière-plan, et l'état de votre agenda.")
    shot(doc, "la fenêtre « Paramètres » avec ses trois sections : Clé API "
              "Mistral, Calendrier Microsoft et Arrière-plan.",
         image="Capture d'écran 2026-05-21 182013.png", width=5.0)

    h2(doc, "Option de confort",
       "Rester actif en arrière-plan et démarrer avec Windows")
    body(doc, "Dans les Paramètres, deux interrupteurs facultatifs :")
    bullet(doc, [("« Continuer en arrière-plan à la fermeture »", True),
                 (" : en fermant la fenêtre, l'application reste discrètement "
                  "active pour vous prévenir de vos prochaines réunions.",
                  False)])
    bullet(doc, [("« Lancer au démarrage Windows »", True),
                 (" : l'application se lance toute seule à l'ouverture de "
                  "votre session, réduite, prête à vous envoyer les rappels.",
                  False)])
    shot(doc, "section « Arrière-plan » dans les Paramètres : les deux "
              "interrupteurs « Continuer en arrière-plan à la fermeture » "
              "et « Lancer au démarrage Windows ».",
         image="Capture d'écran 2026-05-21 182140.png", width=6.0)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 4 — Connecter l'agenda (facultatif)
    # ----------------------------------------------------------------- #
    h1(doc, 4, "Connecter votre agenda Outlook (facultatif)")
    body(doc, "Étape facultative : elle permet de retrouver vos réunions "
              "Outlook directement dans l'application. Vous pouvez l'ignorer "
              "et utiliser l'application sans agenda.",
         italic=True, color=MUTED)

    h2(doc, "Connexion", "Relier votre compte Microsoft")
    bullet(doc, [("Ouvrez l'onglet ", False), ("Agenda", True),
                 (", puis cliquez sur « Connecter mon agenda Microsoft ».",
                  False)])
    bullet(doc, "Un code court s'affiche à l'écran.")
    bullet(doc, [("Cliquez sur ", False), ("Ouvrir la page de connexion", True),
                 (", saisissez le code, puis connectez-vous avec votre "
                  "compte Microsoft habituel.", False)])
    bullet(doc, "Votre mot de passe n'est jamais enregistré ; l'application "
                "lit seulement votre agenda.", level=1)
    shot(doc, "l'écran « Connexion Microsoft » dans l'application : le code "
              "à saisir et le bouton « Ouvrir la page de connexion ».",
         image="Capture d'écran 2026-05-21 181727.png", width=6.0)

    h2(doc, "Résultat", "Vos réunions s'affichent")
    bullet(doc, "Vos réunions du jour et des prochains jours apparaissent "
                "dans l'onglet Agenda.")
    bullet(doc, "Chaque réunion porte une pastille d'état : « À enregistrer », "
                "« En cours » ou « ✓ Compte rendu ».")
    shot(doc, "l'onglet « Agenda » — liste des réunions d'aujourd'hui et "
              "des prochains jours, chacune avec son bouton « À "
              "enregistrer ».",
         image="Capture d'écran 2026-05-21 181923.png", width=6.0)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 5 — Capturer une réunion
    # ----------------------------------------------------------------- #
    h1(doc, 5, "Capturer une réunion")
    body(doc, "Quatre façons de capturer une réunion. Choisissez celle qui "
              "vous convient.", italic=True, color=MUTED)

    h2(doc, "Méthode 1", "Enregistrer une réunion de votre agenda")
    bullet(doc, "Sur l'accueil, cliquez sur la réunion concernée.")
    bullet(doc, [("Cliquez sur ", False), ("Enregistrer", True), (".", False)])
    bullet(doc, "Le titre, les participants et le contexte sont déjà "
                "remplis pour vous.")
    shot(doc, "chaque réunion de l'agenda porte un bouton « À enregistrer » "
              "à droite — c'est lui qui démarre la capture.",
         image="Capture d'écran 2026-05-21 181923.png", width=6.0)

    h2(doc, "Méthode 2", "Enregistrer une réunion non planifiée (en direct)")
    bullet(doc, [("Cliquez sur ", False), ("Nouvelle réunion", True),
                 (" (ou « Hors agenda »).", False)])
    bullet(doc, [("Sur l'écran « Capturez votre réunion », gardez l'onglet ",
                  False), ("Enregistrer", True), (".", False)])
    bullet(doc, [("Cliquez sur ", False), ("Démarrer", True),
                 (" : un minuteur défile pendant l'enregistrement.", False)])
    bullet(doc, [("À la fin, cliquez sur ", False), ("Arrêter", True),
                 (".", False)])
    tip(doc, "Vous pouvez réduire la fenêtre pendant l'enregistrement : une "
             "notification Windows vous préviendra dès que le compte rendu "
             "est prêt.")
    shot(doc, "écran « Capturez votre réunion », onglet « Enregistrer » : "
              "le minuteur et le bouton rouge « Démarrer ».",
         image="Capture d'écran 2026-05-21 181939.png", width=6.0)

    h2(doc, "Méthode 3", "Importer un fichier audio existant")
    bullet(doc, [("Cliquez sur « Nouvelle réunion » puis sur l'onglet ",
                  False), ("Audio", True), (".", False)])
    bullet(doc, "Glissez votre fichier audio dans la zone prévue "
                "(« Glissez votre audio ici »), ou cliquez pour le choisir.")
    bullet(doc, "Les formats audio courants et les vidéos MP4 sont acceptés.")
    shot(doc, "onglet « Audio » avec la zone « Glissez votre audio ici · "
              "ou cliquez pour parcourir ».",
         image="Capture d'écran 2026-05-21 181947.png", width=6.0)

    h2(doc, "Méthode 4", "Importer un transcript déjà écrit")
    bullet(doc, [("Cliquez sur « Nouvelle réunion » puis sur l'onglet ",
                  False), ("Transcript", True), (".", False)])
    bullet(doc, "Glissez un fichier .docx (transcript Teams) ou un fichier "
                ".txt dans la zone prévue.")
    bullet(doc, "L'application génère directement le compte rendu à partir "
                "de ce texte.")
    shot(doc, "onglet « Transcript » avec la zone « Glissez votre "
              "transcript ici » (accepte .docx Teams ou .txt).",
         image="Capture d'écran 2026-05-21 181953.png", width=6.0)

    h2(doc, "Préparer", "Renseigner le contexte de la réunion")
    bullet(doc, [("Contexte", True),
                 (" : le sujet, les enjeux, les décisions attendues.",
                  False)])
    bullet(doc, [("Participants", True),
                 (" : les noms, séparés par des virgules.", False)])
    bullet(doc, [("Entreprises", True),
                 (" : les organisations concernées.", False)])
    bullet(doc, "Pour une réunion de l'agenda, ces champs sont déjà "
                "pré-remplis — vérifiez et ajustez si besoin.")
    tip(doc, "Plus le contexte est précis, plus le compte rendu est juste "
             "et pertinent.")
    shot(doc, "le formulaire de contexte avec les trois champs : Contexte, "
              "Participants, Entreprises.")

    h2(doc, "Lancer", "Choisir le mode et lancer le traitement")
    bullet(doc, "Choisissez le mode de génération du compte rendu :")
    bullet(doc, [("Sur cet ordinateur", True),
                 (" (sans connexion Internet, gratuit) — le mode par défaut.",
                  False)], level=1)
    bullet(doc, [("En ligne (Mistral)", True),
                 (" (plus rapide, nécessite une clé à renseigner dans les "
                  "Paramètres).", False)], level=1)
    bullet(doc, [("Cliquez sur ", False), ("Lancer le traitement", True),
                 (".", False)])
    shot(doc, "le bloc de fin du formulaire : choix « Sur cet ordinateur » / "
              "« En ligne (Mistral) » et le bouton « Lancer le traitement ».")

    h2(doc, "Patienter", "Suivre l'avancement")
    bullet(doc, "Une barre de progression montre les étapes : "
                "Conversion → Diarisation → Transcription → Compte rendu.")
    bullet(doc, "L'étape en cours est mise en évidence.")
    bullet(doc, "Vous pouvez continuer à travailler pendant ce temps.")
    bullet(doc, "Comptez environ un cinquième à un tiers de la durée de "
                "la réunion.")
    shot(doc, "la barre de progression du traitement (les étapes Conversion / "
              "Diarisation / Transcription / Compte rendu).")

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 6 — Exploiter le compte rendu
    # ----------------------------------------------------------------- #
    h1(doc, 6, "Exploiter votre compte rendu")

    h2(doc, "Le résultat", "Lire et modifier le compte rendu")
    bullet(doc, "Le compte rendu s'affiche dès qu'il est prêt.")
    bullet(doc, "Vous pouvez le modifier directement dans l'application.")
    bullet(doc, "Tout est enregistré automatiquement (mention "
                "« Enregistré à HH:MM »).")
    bullet(doc, [("Recherchez un mot dans le compte rendu avec ", False),
                 ("Ctrl + F", True), (".", False)])
    shot(doc, "un compte rendu affiché dans l'application, avec son "
              "transcript par tours à droite et la mention "
              "« Auto-enregistrement actif » en haut.",
         image="Capture d'écran 2026-05-21 184506.png", width=6.0)

    h2(doc, "Vérifier", "Consulter le transcript et nommer les intervenants")
    bullet(doc, [("Cliquez sur ", False), ("+ Transcript", True),
                 (" pour afficher le texte complet, ligne par ligne.",
                  False)])
    bullet(doc, "Cliquez sur le nom d'un intervenant pour le remplacer par "
                "un vrai nom (ex. « Intervenant 1 » → « Alice Dupont »).")
    shot(doc, "l'onglet « Transcript » ouvert, avec le menu qui apparaît en "
              "cliquant sur un nom d'intervenant.")

    h2(doc, "Écouter", "Réécouter et télécharger l'audio")
    bullet(doc, [("Cliquez sur ", False), ("Écouter l'enregistrement", True),
                 (" pour ouvrir un petit lecteur audio.", False)])
    bullet(doc, "Un bouton permet de télécharger le fichier audio d'origine.")
    shot(doc, "le compte rendu reste affiché et un petit lecteur audio "
              "apparaît en bas pour réécouter la réunion (avec le bouton "
              "de téléchargement à droite).",
         image="Capture d'écran 2026-05-19 111911.png", width=6.0)

    h2(doc, "Retrouver", "Où sont enregistrés vos fichiers")
    bullet(doc, [("Tout est rangé dans ", False),
                 ("Documents ▸ Réunions", True),
                 (" (un dossier par réunion).", False)])
    bullet(doc, "compte_rendu.docx — le compte rendu Word.", level=1)
    bullet(doc, "transcript.txt — la transcription complète.", level=1)
    bullet(doc, "le fichier audio de la réunion.", level=1)
    bullet(doc, "Vous gérez ces fichiers comme n'importe quel document "
                "(copier, envoyer par e-mail…).")
    bullet(doc, [("Astuce : le bouton ", False),
                 ("Ouvrir le dossier de la réunion", True),
                 (" ouvre directement ce dossier.", False)])
    shot(doc, "l'Explorateur Windows ouvert sur « Documents\\Réunions », "
              "montrant un dossier de réunion et les fichiers qu'il contient.")

    h2(doc, "S'organiser", "Classer et retrouver vos réunions")
    bullet(doc, [("Comptes rendus", True),
                 (" : la liste de toutes vos réunions et leur état.",
                  False)])
    bullet(doc, [("Dossiers", True),
                 (" : créez des dossiers (client, projet…) et déplacez-y "
                  "vos réunions.", False)])
    bullet(doc, [("Recherche", True),
                 (" (Ctrl + F) et mini-calendrier pour filtrer rapidement.",
                  False)])
    bullet(doc, [("Sur chaque réunion : ", False),
                 ("Renommer", True), (", ", False), ("Déplacer", True),
                 (", ", False), ("Supprimer", True), (".", False)])
    shot(doc, "la liste « Comptes rendus » : barre de recherche, filtre "
              "par date, dossiers, et à droite d'une réunion les boutons "
              "Déplacer (dossier) / Renommer (crayon) / Supprimer "
              "(corbeille).",
         image="Capture d'écran 2026-05-19 111726.png", width=3.5)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 7 — Arrière-plan & notifications
    # ----------------------------------------------------------------- #
    h1(doc, 7, "L'application travaille pour vous")

    h2(doc, "Discrète", "Une icône dans la barre des tâches")
    bullet(doc, "Quand vous fermez la fenêtre, l'application reste active en "
                "arrière-plan (si l'option est cochée).")
    bullet(doc, "Une petite icône reste près de l'horloge, en bas à droite.")
    bullet(doc, "Cliquez dessus pour un aperçu rapide ; clic droit pour "
                "« Démarrer un enregistrement » ou « Quitter ».")
    shot(doc, "l'aperçu rapide qui s'ouvre en cliquant sur l'icône Meeting "
              "Assistant dans la barre des tâches.",
         image="Capture d'écran 2026-05-21 184154.png", width=3.5)

    h2(doc, "Rappels", "Des notifications au bon moment")
    bullet(doc, "5 minutes avant une réunion de l'agenda : un rappel pour "
                "penser à l'enregistrer.")
    bullet(doc, "À la fin de la réunion : un rappel pour arrêter "
                "l'enregistrement.")
    bullet(doc, "Quand le compte rendu est prêt : une notification — "
                "cliquez dessus pour l'ouvrir.")
    shot(doc, "exemple : pendant un enregistrement, l'aperçu rapide affiche "
              "le minuteur et le bouton « Arrêter et générer le compte "
              "rendu » ; en bas, une notification Windows confirme que "
              "l'application reste active en arrière-plan.",
         image="Capture d'écran 2026-05-21 184235.png", width=4.0)

    page_break(doc)

    # ----------------------------------------------------------------- #
    #  Partie 8 — Mises à jour et aide
    # ----------------------------------------------------------------- #
    h1(doc, 8, "Mises à jour et aide")

    h2(doc, "Mises à jour", "L'application se met à jour toute seule")
    bullet(doc, "Quand une nouvelle version existe, un message apparaît : "
                "« Une nouvelle version de Meeting Assistant est prête. »")
    bullet(doc, [("Cliquez sur ", False), ("Redémarrer maintenant", True),
                 (" pour l'installer tout de suite, ou sur « Plus tard » "
                  "pour l'installer à la prochaine ouverture.", False)])
    bullet(doc, "Vous n'avez aucune manipulation technique à faire.")
    shot(doc, "la fenêtre « Une nouvelle version est prête » avec les boutons "
              "« Redémarrer maintenant » et « Plus tard ».")

    h2(doc, "Dépannage", "Questions fréquentes")
    faq_table(doc, [
        ("La première ouverture est très longue",
         "C'est normal : l'application télécharge ses outils une seule fois "
         "(environ 2,3 Go). Patientez et ne fermez pas la fenêtre."),
        ("Windows affiche un écran bleu au lancement",
         "Cliquez sur « Informations complémentaires » puis "
         "« Exécuter quand même ». Le fichier n'est pas dangereux."),
        ("Je n'ai pas Internet",
         "Après la première ouverture, l'application fonctionne sans Internet "
         "(mode « Sur cet ordinateur »)."),
        ("Le microphone ne fonctionne pas",
         "Autorisez le microphone dans les réglages de Windows et vérifiez "
         "que le bon micro est sélectionné."),
        ("Le traitement est long",
         "Comptez environ un cinquième à un tiers de la durée de la réunion. "
         "Vous pouvez travailler pendant ce temps."),
        ("Le compte rendu en ligne ne marche pas",
         "Renseignez la clé dans les Paramètres, ou utilisez le mode "
         "« Sur cet ordinateur » qui fonctionne sans clé."),
    ])

    h2(doc, "Bonnes pratiques", "Pour un compte rendu de qualité")
    bullet(doc, "Utilisez un bon microphone et limitez le bruit de fond.")
    bullet(doc, "Renseignez le contexte et les participants : le résultat "
                "est nettement meilleur.")
    bullet(doc, "Relisez et ajustez le compte rendu (il est modifiable).")
    bullet(doc, "Besoin de rapidité ? Utilisez le mode « En ligne (Mistral) ».")
    tip(doc, "Le compte rendu est une aide à la rédaction : une relecture "
             "humaine reste recommandée avant de le diffuser.")

    # Démarrage rapide (récap final)
    p = para(doc, space_before=18, space_after=8)
    _shade_paragraph(p, "EFF4FB")
    _border_paragraph(p, "1E40AF", size=8, space=8)
    run(p, "Démarrage rapide", 14, BRAND, bold=True)
    for s in [
        "1.  Installez l'application.",
        "2.  À la première ouverture, laissez télécharger les outils "
        "(une seule fois).",
        "3.  Cliquez « Nouvelle réunion », ou choisissez une réunion de "
        "votre agenda.",
        "4.  Enregistrez (ou importez l'audio), renseignez le contexte.",
        "5.  Cliquez « Lancer le traitement » et récupérez votre compte "
        "rendu Word.",
    ]:
        bullet(doc, s)

    # Pied de page
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(fp, "Meeting Assistant — Guide d'installation et d'utilisation",
        9, MUTED)

    # Enregistrement
    here = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(here, "Guide_Utilisateur_Meeting_Assistant.docx")
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) // 1024
    print(f"OK — document écrit : {out_path} ({size_kb} Ko)")


if __name__ == "__main__":
    build()
