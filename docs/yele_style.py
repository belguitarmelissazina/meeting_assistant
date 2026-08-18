"""Socle de mise en forme Word — charte « Yele Consulting / Documentation Technique ».

Reprend l'esthetique du corpus Equipe GenAI :
  - bandeau d'en-tete grenat + filet, pied de page « Page X / Y »
  - titres de section grenat avec filet, sous-titres noirs, intertitres grenat
  - tables « cle/valeur » (colonne gauche rose, libelle grenat)
  - tables pleines (bandeau d'en-tete grenat, texte blanc)
  - encadres de note (fond rose, filet grenat a gauche, italique)
  - blocs de code (fond sombre, texte clair, chasse fixe)
  - legendes de schema (bandeau rose, texte grenat italique)
  - couvertures de partie (image pleine page generee par yele_schemas)

Usage :
    from yele_style import YeleDoc
    d = YeleDoc("Documentation Technique — Meeting Assistant")
    d.h1("1. Vue d'ensemble")
    d.para("...")
    d.kv_table([("Frontend", "Next.js"), ...])
    d.save("sortie.docx")
"""
from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
#   Palette
# ─────────────────────────────────────────────────────────────────────────────
GRENAT_HEX = "B2544F"        # primaire — titres, filets, bandeaux de table
GRENAT_D_HEX = "8E403C"      # variante foncee
ROSE_HEX = "F8EBEA"          # colonne de libelles, encadres de note
ROSE_ALT_HEX = "FCF5F4"      # lignes alternees
GRIS_L_HEX = "D9D9D9"        # filets discrets
CODE_BG_HEX = "1E1E1E"       # fond des blocs de code
BLANC_HEX = "FFFFFF"

GRENAT = RGBColor(0xB2, 0x54, 0x4F)
GRENAT_D = RGBColor(0x8E, 0x40, 0x3C)
NOIR = RGBColor(0x22, 0x22, 0x22)
GRIS = RGBColor(0x5A, 0x5A, 0x5A)
GRIS_F = RGBColor(0x77, 0x77, 0x77)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
CODE_FG = RGBColor(0xE8, 0xE8, 0xE8)

SERIF = "Calibri"
MONO = "Consolas"

TEXT_W_CM = 16.0             # largeur utile (A4 - marges)


# ─────────────────────────────────────────────────────────────────────────────
#   Helpers bas niveau (XML)
# ─────────────────────────────────────────────────────────────────────────────
def _shade(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=140, right=140) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _cell_valign(cell, val="center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    e = OxmlElement("w:vAlign")
    e.set(qn("w:val"), val)
    tcPr.append(e)


def _table_borders(table, edges: dict) -> None:
    """edges : {edge: (sz_huitiemes_pt, hex)} ; les autres aretes -> nil."""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        if edge in edges:
            sz, color = edges[edge]
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:color"), color)
            e.set(qn("w:space"), "0")
        else:
            e.set(qn("w:val"), "nil")
        b.append(e)
    tblPr.append(b)


def _par_border(paragraph, edge: str, color: str, sz: int, space: int = 4) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bdr = pPr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        pPr.append(bdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:color"), color)
    e.set(qn("w:space"), str(space))
    bdr.append(e)


def _par_shade(paragraph, hex_fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _fixed_layout(table) -> None:
    table.autofit = False
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblLayout"))
    if old is not None:
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _repeat_header(row) -> None:
    """Repete la ligne d'en-tete quand la table franchit une page."""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def _field(paragraph, instr: str):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(it)
    run._r.append(f2)
    return run


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


# ─────────────────────────────────────────────────────────────────────────────
#   Document
# ─────────────────────────────────────────────────────────────────────────────
class YeleDoc:
    """Fabrique de document Word a la charte Yele."""

    def __init__(self, sous_titre: str = "Documentation Technique — Equipe GenAI"):
        self.doc = Document()
        self.sous_titre = sous_titre
        self._setup_page()
        self._setup_base_style()
        self._setup_header()
        self._setup_footer()

    # ── Mise en page ────────────────────────────────────────────────────────
    def _setup_page(self) -> None:
        s = self.doc.sections[0]
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.0)
        s.header_distance = Cm(1.1)
        s.footer_distance = Cm(1.0)

    def _setup_base_style(self) -> None:
        st = self.doc.styles["Normal"]
        st.font.name = SERIF
        st.font.size = Pt(10.5)
        st.font.color.rgb = NOIR
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.15
        # Police de repli pour les caracteres non latins
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        rf.set(qn("w:eastAsia"), SERIF)

    def _setup_header(self) -> None:
        hp = self.doc.sections[0].header.paragraphs[0]
        hp.text = ""
        hp.paragraph_format.space_after = Pt(2)
        r = hp.add_run("Yele Consulting")
        r.bold = True
        r.underline = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRENAT
        r2 = hp.add_run("   |   " + self.sous_titre + "   |   Confidentiel")
        r2.font.size = Pt(9)
        r2.font.color.rgb = GRIS_F
        _par_border(hp, "bottom", GRENAT_HEX, 8, space=3)

    def _setup_footer(self) -> None:
        fp = self.doc.sections[0].footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _par_border(fp, "top", GRIS_L_HEX, 4, space=4)
        r = fp.add_run("Page ")
        r.font.size = Pt(8)
        r.font.color.rgb = GRIS
        pf = _field(fp, "PAGE")
        pf.font.size = Pt(8)
        pf.font.color.rgb = GRIS
        r2 = fp.add_run(" / ")
        r2.font.size = Pt(8)
        r2.font.color.rgb = GRIS
        np = _field(fp, "NUMPAGES")
        np.font.size = Pt(8)
        np.font.color.rgb = GRIS

    # ── Blocs de texte ──────────────────────────────────────────────────────
    def h1(self, txt: str):
        """Titre de section : grenat, gras, filet grenat en dessous."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(17)
        r.font.color.rgb = GRENAT
        _par_border(p, "bottom", GRENAT_HEX, 10, space=5)
        return p

    def h2(self, txt: str):
        """Sous-section : noir, gras."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = NOIR
        return p

    def h3(self, txt: str):
        """Intertitre : grenat, gras, petite taille."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRENAT
        return p

    def h4(self, txt: str):
        """Sous-intertitre : noir, gras, souligne."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.bold = True
        r.underline = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = NOIR
        return p

    def para(self, txt: str, italic: bool = False, size: float = 10.5,
             color: RGBColor | None = None, align=None):
        """Paragraphe. Le balisage **gras** est interprete."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        if align is not None:
            p.alignment = align
        self._runs(p, txt, italic=italic, size=size, color=color or NOIR)
        return p

    def _runs(self, p, txt: str, italic=False, size=10.5,
              color=NOIR, bold=False, mono=False):
        """Ecrit `txt` en gerant **gras**, `code` et le retour a la ligne."""
        import re
        for chunk in txt.split("\n"):
            if chunk is not txt.split("\n")[0]:
                p.add_run().add_break()
            for part in re.split(r"(\*\*.+?\*\*|`.+?`)", chunk):
                if not part:
                    continue
                r = p.add_run()
                if part.startswith("**") and part.endswith("**"):
                    r.text = part[2:-2]
                    r.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    r.text = part[1:-1]
                    r.font.name = MONO
                    r.font.size = Pt(size - 0.7)
                    r.font.color.rgb = GRENAT_D
                    continue
                else:
                    r.text = part
                    r.bold = bold
                r.italic = italic
                r.font.size = Pt(size)
                r.font.color.rgb = color
                if mono:
                    r.font.name = MONO

    def bullets(self, items: list[str], size: float = 10.5):
        """Liste a puces ● comme dans le corpus de reference."""
        for it in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run("● ")
            r.font.size = Pt(7)
            r.font.color.rgb = GRENAT
            self._runs(p, it, size=size)
        self.spacer(4)

    def numbered(self, items: list[str], size: float = 10.5):
        for i, it in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.9)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{i}. ")
            r.bold = True
            r.font.size = Pt(size)
            r.font.color.rgb = GRENAT
            self._runs(p, it, size=size)
        self.spacer(4)

    def spacer(self, pt: int = 8):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)
        for r in p.runs:
            r.font.size = Pt(1)
        return p

    # ── Encadres ────────────────────────────────────────────────────────────
    def note(self, txt: str, accent: str = GRENAT_HEX, fill: str = ROSE_HEX):
        """Encadre d'information : fond rose, filet grenat a gauche, italique."""
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_layout(t)
        _table_borders(t, {"left": (18, accent)})
        c = t.cell(0, 0)
        c.width = Cm(TEXT_W_CM)
        _shade(c, fill)
        _cell_margins(c, top=110, bottom=110, left=180, right=160)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        self._runs(p, txt, italic=True, size=9.5, color=GRENAT_D)
        self.spacer(9)
        return t

    def warning(self, txt: str):
        """Encadre d'alerte — meme forme, accent fonce."""
        return self.note("⚠  " + txt, accent=GRENAT_D_HEX, fill="FBEDEC")

    def code(self, txt: str, size: float = 8.5):
        """Bloc de code : fond sombre, texte clair, chasse fixe."""
        lines = txt.strip("\n").split("\n")
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_layout(t)
        _table_borders(t, {})
        c = t.cell(0, 0)
        c.width = Cm(TEXT_W_CM)
        _shade(c, CODE_BG_HEX)
        _cell_margins(c, top=110, bottom=110, left=160, right=140)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for i, ln in enumerate(lines):
            if i:
                p.add_run().add_break()
            r = p.add_run(ln if ln else " ")
            r.font.name = MONO
            r.font.size = Pt(size)
            r.font.color.rgb = CODE_FG
        self.spacer(9)
        return t

    # ── Tables ──────────────────────────────────────────────────────────────
    def kv_table(self, rows: list[tuple[str, str]], label_w: float = 5.2):
        """Table « cle / valeur » : colonne gauche rose + libelle grenat."""
        t = self.doc.add_table(rows=len(rows), cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_layout(t)
        _table_borders(t, {"insideH": (4, GRIS_L_HEX), "bottom": (4, GRIS_L_HEX)})
        val_w = TEXT_W_CM - label_w
        for i, (k, v) in enumerate(rows):
            cl, cv = t.rows[i].cells
            cl.width = Cm(label_w)
            cv.width = Cm(val_w)
            _shade(cl, ROSE_HEX)
            for c in (cl, cv):
                _cell_margins(c)
                _cell_valign(c)
            pl = cl.paragraphs[0]
            pl.paragraph_format.space_after = Pt(0)
            rl = pl.add_run(k)
            rl.bold = True
            rl.font.size = Pt(9.5)
            rl.font.color.rgb = GRENAT
            pv = cv.paragraphs[0]
            pv.paragraph_format.space_after = Pt(0)
            self._runs(pv, v, size=9.5)
        self.spacer(10)
        return t

    def table(self, headers: list[str], rows: list[list[str]],
              widths: list[float] | None = None, zebra: bool = True):
        """Table pleine : bandeau d'en-tete grenat, texte blanc."""
        n = len(headers)
        if widths is None:
            widths = [TEXT_W_CM / n] * n
        else:
            s = sum(widths)
            widths = [w * TEXT_W_CM / s for w in widths]

        t = self.doc.add_table(rows=len(rows) + 1, cols=n)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_layout(t)
        _table_borders(t, {"insideH": (4, GRIS_L_HEX), "bottom": (4, GRIS_L_HEX)})

        hdr = t.rows[0]
        _repeat_header(hdr)
        for j, h in enumerate(headers):
            c = hdr.cells[j]
            c.width = Cm(widths[j])
            _shade(c, GRENAT_HEX)
            _cell_margins(c, top=70, bottom=70)
            _cell_valign(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = BLANC

        for i, row in enumerate(rows):
            cells = t.rows[i + 1].cells
            for j, val in enumerate(row):
                c = cells[j]
                c.width = Cm(widths[j])
                if zebra and i % 2 == 1:
                    _shade(c, ROSE_ALT_HEX)
                _cell_margins(c)
                _cell_valign(c, "top")
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                self._runs(p, str(val), size=9)
        self.spacer(10)
        return t

    # ── Schemas ─────────────────────────────────────────────────────────────
    def schema_caption(self, txt: str):
        """Bandeau de legende rose, texte grenat italique — avant l'image."""
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_layout(t)
        _table_borders(t, {"top": (4, GRENAT_HEX), "bottom": (4, GRENAT_HEX),
                           "left": (4, GRENAT_HEX), "right": (4, GRENAT_HEX)})
        c = t.cell(0, 0)
        c.width = Cm(TEXT_W_CM)
        _shade(c, ROSE_HEX)
        _cell_margins(c, top=70, bottom=70, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt)
        r.bold = True
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRENAT
        self.spacer(6)
        return t

    def schema(self, png_path: str, caption: str | None = None,
               width_cm: float = TEXT_W_CM):
        """Legende + image centree."""
        if caption:
            self.schema_caption(caption)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(png_path), width=Cm(width_cm))
        return p

    def cover(self, png_path: str):
        """Couverture de partie : image pleine largeur."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(png_path), width=Cm(TEXT_W_CM))
        return p

    # ── Structure ───────────────────────────────────────────────────────────
    def page_break(self):
        self.doc.add_page_break()

    def toc_entry(self, num: str, titre: str, page_hint: str = ""):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(TEXT_W_CM),
                                                  WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run(num + "  ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = GRENAT
        r2 = p.add_run(titre)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = NOIR
        if page_hint:
            r3 = p.add_run("\t" + page_hint)
            r3.font.size = Pt(9.5)
            r3.font.color.rgb = GRIS
        return p

    def save(self, path: str) -> str:
        self.doc.save(path)
        return path
