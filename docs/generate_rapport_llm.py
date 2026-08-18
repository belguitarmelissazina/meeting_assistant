"""Genere le rapport Word de la PARTIE LLM (mise en page soignee).

Structure en 3 parties : 1) chunking, 2) architectures (extraction puis
agentique), 3) optimisations. Chaque technique a une analyse "pourquoi ca
marche / pourquoi pas" + des sources academiques verifiees (liens cliquables).
Toutes les comparaisons portent sur LA MEME reunion (RTE : Jerome Picault /
Maya Sahraoui). Chiffres verbatim des metrics.json / logs.

Style : boites colorees (vert = adopte, rouge = rejete, bleu = analyse,
jaune = note), tableaux a en-tete bleu + zebrage, nombres surlignes vert/rouge.
Aucun em-dash ni fleche dans le texte (consigne).
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette (sobre, "cabinet de conseil" : 1 primaire bleu ardoise + 1 accent cuivre)
BLEU = RGBColor(0x1F, 0x3A, 0x5F)    # primaire (titres h1, hyperliens, regle d'en-tete)
BLEU2 = RGBColor(0x2E, 0x4D, 0x6B)   # h2, captions, liens
GRIS = RGBColor(0x5A, 0x5A, 0x5A)    # legendes, h3
NOIR = RGBColor(0x22, 0x22, 0x22)    # corps de texte
VERT = RGBColor(0x2E, 0x6B, 0x3F)    # semantique positif
ROUGE = RGBColor(0x9E, 0x2B, 0x25)   # semantique negatif
ACCENT = RGBColor(0xB5, 0x65, 0x1D)  # cuivre (filet sous h1, couverture)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
# Fonds tres pales (hex sans #) et filets
ACCENT_HEX = "B5651D"
RULE = "BFBFBF"        # filets horizontaux de tableau
HDR_RULE = "1F3A5F"    # filet sous l'en-tete de tableau
VERT_BG = "EAF1EC"
ROUGE_BG = "F7ECEB"
BLEU_BG = "EEF2F6"
JAUNE_BG = "FBF4EA"
GRIS_BG = "F4F5F7"
HDR_BG = "1F3A5F"      # (conserve pour compat, non utilise par le nouveau style)
ZEBRA = "F7F8FA"
CONTENT_IN = 6.06      # largeur de la colonne de texte en pouces (~15.4 cm)

doc = Document()
_sec = doc.sections[0]
_sec.page_width = Mm(210)
_sec.page_height = Mm(297)
_sec.top_margin = Cm(2.5)
_sec.bottom_margin = Cm(2.5)
_sec.left_margin = Cm(2.8)
_sec.right_margin = Cm(2.8)
_sec.footer_distance = Cm(1.2)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = NOIR
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)


# --------------------------------------------------------------------------
# Helpers bas niveau
# --------------------------------------------------------------------------
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_margins(cell, top=50, bottom=50, left=120, right=120):
    """Marges internes de cellule en twips (1/20 pt) pour aerer."""
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _hide_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tcPr.append(b)


def _cell_border(cell, edge, color, sz):
    """sz en huitiemes de point (8 = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders")
        tcPr.append(b)
    e = b.find(qn(f"w:{edge}"))
    if e is None:
        e = OxmlElement(f"w:{edge}")
        b.append(e)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:color"), color)
    e.set(qn("w:space"), "0")


def _table_borders(table, edges):
    """edges : dict {edge: (sz, color)} ; les autres aretes -> nil (invisibles)."""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        if edge in edges:
            sz, color = edges[edge]
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(sz))
            e.set(qn("w:color"), color)
            e.set(qn("w:space"), "0")
        else:
            e.set(qn("w:val"), "nil")
        b.append(e)
    tblPr.append(b)


def _par_border(paragraph, edge, color, sz, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    bdr = pPr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        pPr.append(bdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:color"), color)
    e.set(qn("w:space"), str(space))
    bdr.append(e)


def _fixed_layout(table):
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblLayout"))
    if old is not None:
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _page_field(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(it)
    run._r.append(f2)
    return run


def setup_footer():
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(Cm(15.4), WD_TAB_ALIGNMENT.RIGHT)
    _par_border(fp, "top", "D9D9D9", 4, space=4)
    r = fp.add_run("Yele Consulting · Rapport technique (partie LLM)")
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS
    t = fp.add_run("\t")
    t.font.size = Pt(8)
    r2 = fp.add_run("Page ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRIS
    pf = _page_field(fp)
    pf.font.size = Pt(8)
    pf.font.color.rgb = GRIS


def _spacer(pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.space_before = Pt(0)
    return p


def _no_table_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _cell(cell, text, size=9, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return r


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


# --------------------------------------------------------------------------
# Helpers de contenu
# --------------------------------------------------------------------------
def _heading(level, txt, font_pt, color, sb, sa, rule=False):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(txt)
    r.font.name = "Georgia"
    r.font.size = Pt(font_pt)
    r.bold = True
    r.font.color.rgb = color
    if rule:
        _par_border(p, "bottom", ACCENT_HEX, 12, space=6)  # filet cuivre 1.5 pt sous h1
    return p


def h1(txt):
    return _heading(1, txt, 18, BLEU, 18, 6, rule=True)


def h2(txt):
    return _heading(2, txt, 13, BLEU2, 14, 4)


def h3(txt):
    return _heading(3, txt, 11, GRIS, 10, 3)


def para(txt, italic=False, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(txt, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(txt)
    else:
        p.add_run(txt)
    return p


def _box(label, text, fill, label_color):
    """Callout = simple rectangle (fond tres pale + fine bordure de la couleur)."""
    hexc = str(label_color)
    t = doc.add_table(rows=1, cols=1)
    _fixed_layout(t)
    _table_borders(t, {
        "top": (6, hexc), "bottom": (6, hexc),
        "left": (6, hexc), "right": (6, hexc),
    })
    c = t.rows[0].cells[0]
    c.width = Cm(15.4)
    c.text = ""
    _shade(c, fill)
    _cell_margins(c, top=80, bottom=80, left=160, right=160)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if label:
        r = p.add_run(label + "   ")
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = label_color
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = NOIR
    _spacer(8)
    return t


def verdict(txt):
    """Boite verte (adopte), rouge (rejete) ou bleue (neutre) selon le contenu."""
    low = txt.lower()
    if low.startswith(("adopt", "retenu", "prometteur")):
        _box("Verdict", txt, VERT_BG, VERT)
    elif low.startswith(("rejet", "ecart", "écart", "non concl", "abandonn")):
        _box("Verdict", txt, ROUGE_BG, ROUGE)
    else:
        _box("Verdict", txt, BLEU_BG, BLEU)


def pourquoi(txt):
    _box("Analyse", txt, BLEU_BG, BLEU)


def note(txt):
    _box("Note", txt, JAUNE_BG, RGBColor(0x8A, 0x6D, 0x00))


def table(headers, rows, widths=None, green=None, red=None):
    """Tableau a en-tete bleu, zebrage, et surlignage vert/rouge.

    green / red : ensembles de (i_ligne_corps, j_colonne) a colorer.
    Auto : '-NN %' -> vert ; mots negatifs -> rouge ; 'ADOPTE/retenu' -> vert.
    """
    green = green or set()
    red = red or set()
    NEG = ("plus lent", "plus lente", "incomplet", "plante", "vide", "rejet")
    POS = ("adopt", "retenu")
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed_layout(t)
    # Filets horizontaux uniquement : fines regles grises entre lignes + bas.
    _table_borders(t, {"insideH": (4, RULE), "bottom": (6, RULE)})
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]
        _cell(c, htext, size=9, bold=True, color=BLEU)   # en-tete blanc, texte primaire
        _cell_margins(c)
        _cell_border(c, "bottom", HDR_RULE, 12)          # regle primaire 1.5 pt sous l'en-tete
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            sval = str(val)
            low = sval.lower()
            color = None
            bold = False
            if (i, j) in green:
                color, bold = VERT, True
            elif (i, j) in red:
                color, bold = ROUGE, True
            elif sval.startswith("-") and "%" in sval:
                color, bold = VERT, True
            elif any(k in low for k in NEG):
                color, bold = ROUGE, True
            elif any(low.startswith(k) for k in POS):
                color, bold = VERT, True
            _cell(cells[j], sval, size=9, bold=bold, color=color)
            _cell_margins(cells[j])
    # Largeurs : proportions fournies, mises a l'echelle de la colonne de texte.
    if widths:
        tot = sum(widths)
        scaled = [CONTENT_IN * w / tot for w in widths]
    else:
        scaled = [CONTENT_IN / len(headers)] * len(headers)
    for row in t.rows:
        for j, w in enumerate(scaled):
            row.cells[j].width = Inches(w)
    _spacer(8)
    return t


def extrait_cr(source_filename, lignes, intro=None):
    if intro:
        para(intro, italic=True, size=9.5, color=GRIS)
    t = doc.add_table(rows=1, cols=1)
    _fixed_layout(t)
    _table_borders(t, {
        "top": (6, "C9CCD1"), "bottom": (6, "C9CCD1"),
        "left": (6, "C9CCD1"), "right": (6, "C9CCD1"),
    })
    c = t.rows[0].cells[0]
    c.width = Cm(15.4)
    c.text = ""
    _shade(c, GRIS_BG)
    _cell_margins(c, top=80, bottom=80, left=160, right=160)
    for i, ln in enumerate(lignes):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = NOIR
    p = doc.add_paragraph()
    r = p.add_run("Extrait de : " + source_filename)
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = BLEU2
    _spacer(8)


def side_by_side(colA_title, colA_text, colA_src, colB_title, colB_text, colB_src,
                 sentimentA="neutre", sentimentB="neutre"):
    """Deux extraits cote a cote ; en-tete teinte vert/rouge/bleu selon sentiment."""
    fills = {"bon": VERT_BG, "mauvais": ROUGE_BG, "neutre": BLEU_BG}
    txts = {"bon": VERT, "mauvais": ROUGE, "neutre": BLEU}
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed_layout(t)
    _table_borders(t, {"insideV": (4, RULE)})   # seul un fin filet vertical separe les colonnes
    for j, (ti, sent) in enumerate([(colA_title, sentimentA), (colB_title, sentimentB)]):
        c = t.rows[0].cells[j]
        c.width = Cm(7.55)
        _shade(c, fills[sent])
        _cell(c, ti, size=9.5, bold=True, color=txts[sent])
        _cell_margins(c, top=70, bottom=50, left=150, right=150)
    for j, tx in enumerate([colA_text, colB_text]):
        c = t.rows[1].cells[j]
        c.width = Cm(7.55)
        c.text = ""
        _cell_margins(c, top=50, bottom=70, left=150, right=150)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(tx)
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = NOIR
    p = doc.add_paragraph()
    r = p.add_run("Sources : " + colA_src + "   |   " + colB_src)
    r.font.size = Pt(8)
    r.italic = True
    r.font.color.rgb = BLEU2
    _spacer(8)


def sources(items):
    p = doc.add_paragraph()
    r = p.add_run("Source(s) : ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRIS
    for i, (label, url) in enumerate(items):
        if i > 0:
            s = p.add_run("    ")
            s.font.size = Pt(8.5)
        _add_hyperlink(p, url, label)
    doc.add_paragraph()
    return p


# ===========================================================================
# PAGE DE TITRE
# ===========================================================================
setup_footer()

# Filet primaire en haut
top_rule = doc.add_paragraph()
top_rule.paragraph_format.space_after = Pt(0)
_par_border(top_rule, "bottom", "1F3A5F", 18, space=1)

# Espace pour centrer visuellement le bloc titre
for _ in range(4):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Sur-titre (eyebrow)
eyebrow = doc.add_paragraph()
eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
eyebrow.paragraph_format.space_after = Pt(10)
r = eyebrow.add_run("RAPPORT TECHNIQUE · PARTIE LLM")
r.font.size = Pt(10.5)
r.font.color.rgb = ACCENT
r.bold = True

# Titre
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.line_spacing = 1.1
title.paragraph_format.space_after = Pt(8)
r = title.add_run("Génération de comptes rendus de réunion par LLM local")
r.bold = True
r.font.name = "Georgia"
r.font.size = Pt(26)
r.font.color.rgb = BLEU

# Sous-titre
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Du découpage du transcript à l'architecture agentique : "
                "architectures, optimisations et benchmarks des modèles")
r.font.name = "Georgia"
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GRIS

# Grand espace
for _ in range(6):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Bloc meta
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Yele Consulting, rapport technique (partie LLM)\n"
                 "Période des expériences : mars à juin 2026")
r.font.size = Pt(10.5)
r.font.color.rgb = GRIS

# Filet cuivre en bas
bot_rule = doc.add_paragraph()
bot_rule.paragraph_format.space_before = Pt(8)
_par_border(bot_rule, "bottom", ACCENT_HEX, 8, space=1)

doc.add_page_break()

# ===========================================================================
# 0. CONTEXTE
# ===========================================================================
h1("0. Contexte, objectif et méthode")
para("Objectif du module LLM : transformer le transcript d'une réunion (issu de la "
     "diarisation et de la transcription locales) en un compte rendu structuré "
     "(synthèse, sujets, décisions, plan d'action), avec trois contraintes non "
     "négociables :")
bullet("100 % local, sur CPU uniquement (cible 8 à 16 Go de RAM), sans GPU ni API cloud.", "Contrainte 1 : ")
bullet("Zéro hallucination : ne restituer que ce qui est explicitement dit.", "Contrainte 2 : ")
bullet("Fonctionner aussi bien sur un enregistrement live que sur un audio importé.", "Contrainte 3 : ")

para("Stack commune à toutes les expériences : llama-server.exe (llama.cpp, API "
     "compatible OpenAI) pour l'inférence des modèles GGUF quantifiés ; embeddings "
     "all-MiniLM-L6-v2 (384 dim, CPU) pour le découpage sémantique ; sortie JSON "
     "contrainte via response_format et grammaire GBNF.")

para("La réunion de référence (réunion RTE)", bold=True, color=BLEU)
para("Sauf mention contraire, toutes les expériences et tous les extraits de comptes "
     "rendus de ce rapport portent sur LA MÊME réunion : un échange d'environ 1 h, "
     "environ 9 participants (dont Jérôme Picault et Maya Sahraoui), autour de l'IA "
     "générative et de l'orchestration d'outils (MCP) pour RTE. Elle est déclinée en "
     "plusieurs transcripts : transcript1.txt (556 segments), transcript_formatted.txt "
     "(247 seg) et dicte_audio_3.normalized.txt (266 seg), ce dernier servant de "
     "référence pour la majorité des benchmarks. Comparer toujours la même réunion "
     "permet de juger l'effet d'une seule variable à la fois (découpage, modèle ou "
     "optimisation).")

note("Il n'existe pas de gold standard humain dans le projet. Les scores de qualité "
     "automatiques (BERTScore, ROUGE, LLM-as-judge) ont été codés mais jamais exécutés "
     "sur ces runs. Les comparaisons portent donc sur des métriques OBJECTIVES (durée, "
     "RAM, nombre d'appels LLM, nombre d'items extraits) et sur une évaluation "
     "QUALITATIVE manuelle. Tout chiffre cité provient verbatim des fichiers "
     "metrics.json et des logs des benchmarks. Les comptes rendus complets cités en "
     "extrait sont rassemblés dans docs/comptes_rendus_references/ (un fichier par "
     "technique, architecture ou modèle).")

doc.add_page_break()

# ===========================================================================
# PARTIE 1 — CHUNKING
# ===========================================================================
h1("Partie 1.  Le découpage du transcript (chunking)")
para("Un transcript d'1 h dépasse le contexte utile d'un petit modèle CPU et le "
     "noierait dans le bruit. Il faut donc le découper en morceaux (chunks) traités un "
     "par un. Tout le pipeline dépend de la qualité de ce découpage : un mauvais "
     "découpage donne des sujets coupés en deux ou mélangés. Plusieurs approches ont "
     "été testées (dossier benchmark_llm, mars et avril 2026).", italic=True)

h2("1.1  Évolution des approches de découpage")
para("Trois familles successives, chacune corrigeant un défaut de la précédente. "
     "Chaque ligne de tableau est une expérience distincte, sur la réunion RTE.")

h3("Approche 1 : clustering thématique HDBSCAN (V1)")
para("Fenêtres glissantes (4/2), embeddings MiniLM, clustering HDBSCAN, 1 appel LLM "
     "par cluster, assemblage LLM final.")
bullet("Défaut constaté : clusters non monotones dans le temps, 30 à 40 % du contenu jeté comme bruit.")
verdict("Abandonné. Le clustering perd l'ordre chronologique de la réunion.")

h3("Approche 2 : HDBSCAN amélioré et post-traitement (V2)")
para("Fenêtres 8/3, HDBSCAN(min=5) et fusion de centroïdes, réassignation du bruit, "
     "split KMeans des gros clusters, résolution des locuteurs par regex, prompts "
     "anti-hallucination durcis. Mesuré sur transcript1.txt (556 seg).")
table(
    ["Modèle", "Durée", "RAM serveur", "Clusters"],
    [
        ["Ministral 3B", "39 min 22", "3 786 Mo", "5"],
        ["Qwen2.5 3B", "20 min 48", "2 670 Mo", "5"],
        ["Qwen3 4B", "48 min 55", "4 670 Mo", "5"],
        ["SmolLM3 3B", "24 min 48", "3 870 Mo", "5"],
    ],
    widths=[2.0, 1.6, 1.6, 1.2],
)
para("CR complets : chunking_v2_hdbscan_{ministral3b, qwen2.5-3b, qwen3-4b, smollm3}.md",
     italic=True, size=8.5, color=BLEU2)
verdict("Abandonné à terme. Plus robuste que V1 mais toujours dépendant du clustering, et trop lent.")

h3("Approche 3 : détection de frontières sémantiques (boundary detection, V3, RETENUE)")
para("Rupture d'algorithme : on abandonne le clustering au profit d'une détection de "
     "frontières par chute de similarité cosine entre fenêtres consécutives (lissage "
     "gaussien sigma=2, vallées au percentile 5). On obtient des chunks CHRONOLOGIQUES, "
     "puis 1 LLM par chunk, puis assemblage. Mesuré sur transcript1.txt (556 seg, 9 chunks).")
table(
    ["Modèle", "V2 (HDBSCAN)", "V3 (boundary)", "Gain"],
    [
        ["Ministral 3B", "39 min 22", "26 min 37", "-32 %"],
        ["Qwen2.5 3B", "20 min 48", "12 min 13", "-41 %"],
        ["Qwen3 4B", "48 min 55", "32 min 40", "-33 %"],
        ["SmolLM3 3B", "24 min 48", "14 min 56", "-40 %"],
    ],
    widths=[1.9, 1.7, 1.7, 1.1],
)
para("CR complets : chunking_v3_boundary_{ministral3b, qwen2.5-3b, qwen3-4b, smollm3}.md",
     italic=True, size=8.5, color=BLEU2)
verdict("Adopté. Le découpage chronologique est plus rapide ET préserve l'ordre du récit. C'est la base de tout le pipeline suivant.")
pourquoi("HDBSCAN regroupe les passages par densité dans l'espace des embeddings sans "
         "tenir compte du temps : il produit des clusters non contigus (un même thème "
         "revient à 3 moments différents fusionnés) et étiquette comme bruit les points "
         "isolés, d'où les 30 à 40 % de contenu jeté. La détection de frontières fait "
         "l'inverse : elle garde l'ordre linéaire et ne coupe QUE là où la cohésion "
         "lexico-sémantique chute (vallées de similarité cosine entre fenêtres "
         "consécutives). C'est exactement le principe de TextTiling (Hearst, 1997), "
         "transposé ici sur des embeddings de phrases plutôt que sur des comptages de "
         "mots. Résultat : aucun contenu jeté, ordre du récit préservé, et moins de "
         "calcul (pas de clustering itératif).")
sources([
    ("TextTiling, Hearst, Computational Linguistics 1997", "https://aclanthology.org/J97-1003/"),
    ("HDBSCAN, Campello, Moulavi, Sander, PAKDD 2013", "https://link.springer.com/chapter/10.1007/978-3-642-37456-2_14"),
])

h2("1.2  Paramètres et résultats réels du découpage retenu")
para("Implémentation en production (meeting_minutes_pipeline.py) : build_windows, "
     "embed_windows, detect_topic_boundaries, build_topic_chunks, avec re-découpage "
     "récursif (_resplit_semantic) des chunks trop longs au point de plus faible "
     "similarité. Paramètres :")
table(
    ["Paramètre", "Valeur", "Rôle"],
    [
        ["Fenêtre glissante", "3 segments (slide 1)", "granularité de l'analyse"],
        ["Lissage gaussien sigma", "2,0", "atténue le bruit des similarités"],
        ["Percentile de vallée", "5 %", "seuil de détection des frontières"],
        ["Distance min. frontières", "10 fenêtres", "évite les coupures trop proches"],
        ["Taille max. d'un chunk", "15 000 caractères", "déclenche le re-split récursif"],
        ["Modèle d'embeddings", "all-MiniLM-L6-v2 (384d, CPU)", "vectorisation des fenêtres"],
    ],
    widths=[2.1, 2.0, 2.5],
)
para("Sur dicte_audio_3.normalized (266 segments, run ik_llama v5) : 264 fenêtres "
     "glissantes, similarités min 0,444 / moy 0,812 / max 0,967, seuil percentile 5 % = "
     "0,767, soit 6 frontières et 10 chunks chronologiques (de 1 120 à 13 704 "
     "caractères). Coût du découpage : embeddings 11,6 s et construction 10,5 s "
     "(négligeable devant les ~19 min d'appels LLM).")
pourquoi("Le modèle d'embeddings all-MiniLM-L6-v2 est un MiniLM (Wang et al., 2020, "
         "distillation de l'attention d'un grand modèle vers un petit) affiné façon "
         "Sentence-BERT (Reimers et Gurevych, 2019) pour que la similarité cosine entre "
         "deux vecteurs de phrases reflète leur proximité de sens. C'est ce qui rend la "
         "détection de vallées fiable et rapide sur CPU : 384 dimensions, environ 80 Mo, "
         "des centaines de fenêtres encodées en environ 12 s.")
sources([
    ("Sentence-BERT, Reimers et Gurevych, EMNLP 2019", "https://arxiv.org/abs/1908.10084"),
    ("MiniLM, Wang et al., NeurIPS 2020", "https://arxiv.org/abs/2002.10957"),
])

h2("1.3  Approches sans découpage (contre-exemples)")
para("En parallèle, des pipelines sans découpage par frontières ont été testés sur la "
     "même réunion :")
table(
    ["Pipeline", "Principe", "Modèle(s)", "Durée", "Remarque"],
    [
        ["Multi-pass V1", "4 extractions ciblées par chunk et 3 sections (~40 appels)",
         "Mistral 7B", "1 h 46", "lent"],
        ["Multi-pass V1", "idem", "Qwen2.5 3B", "43 min 38", ""],
        ["Multi-pass V1", "idem", "LFM2.5-1.2B", "18 min 14", "rapport quasi vide (295 car.)"],
        ["3-calls", "0 chunking, 3 appels sur transcript entier (ctx 32k)",
         "LFM Extract et Transcript", "6 min 10", "le plus rapide"],
        ["one-shot", "tout le transcript en 1 appel", "LFM2-2.6B", "29 min 09", ""],
        ["one-shot", "tout le transcript en 1 appel", "LFM2.5-1.2B", "4 min 17", "1 509 car. (court)"],
    ],
    widths=[1.3, 2.3, 1.7, 1.0, 1.4],
)
verdict("Les approches sans découpage (3-calls, one-shot) sont spectaculairement rapides "
        "mais produisent des CR trop courts ou incomplets sur un transcript d'1 h. Le "
        "découpage par frontières (approche 3) reste le meilleur compromis "
        "exhaustivité / fiabilité.")

doc.add_page_break()

# ===========================================================================
# PARTIE 2 — ARCHITECTURES
# ===========================================================================
h1("Partie 2.  Les architectures de génération")
para("Une fois le découpage figé (boundary detection), deux générations d'architecture "
     "se sont succédé : d'abord une architecture extraction (assemblage déterministe), "
     "puis une refonte agentique (orchestrateur et workers).", italic=True)

h2("2.A  Architecture extraction et assemblage déterministe")
para("Pour chaque chunk : 1 appel résumé JSON et 1 appel extraction JSON (décisions et "
     "actions, avec few-shot négatifs), sortie contrainte par grammaire. Puis assemblage "
     "100 % Python (zéro LLM), 1 LLM Executive Summary et 1 LLM recommandations. Mesuré "
     "sur dicte_audio_3.normalized (266 seg, 11 chunks).")
table(
    ["Run", "Modèle", "Durée", "RAM pic", "Décisions", "Actions"],
    [
        ["boundary (V3)", "Ministral 3B", "29 min 15", "2 945 Mo", "n.c.", "n.c."],
        ["extraction (V4)", "Ministral 3B", "51 min 52", "5 825 Mo", "0", "0"],
        ["extraction (V4)", "Qwen2.5 3B", "16 min 10", "4 165 Mo", "5", "9"],
    ],
    widths=[1.4, 1.6, 1.4, 1.2, 1.2, 1.2],
    green={(2, 2)}, red={(1, 2)},
)
para("CR complets : archi_extraction_v4_ministral3b.md et archi_extraction_v4_qwen2.5-3b.md",
     italic=True, size=8.5, color=BLEU2)
verdict("Architecture retenue (extraction JSON et assemblage déterministe), pour la "
        "fiabilité du format. Constat fort : sur ce même schéma, Qwen2.5 3B est environ "
        "3 fois plus rapide que Ministral 3B et ose extraire des items.")

h3("Nuance critique : la quantité d'items extraits n'est pas la qualité")
note("Sur cette réunion (une prise de contact), il n'y a quasiment AUCUNE décision "
     "réellement actée. Les 5 décisions extraites par Qwen2.5 3B, vérifiées contre le "
     "transcript, sont en réalité : 2 récits d'un projet PASSÉ de Maya, 1 non-décision "
     "(Aucune décision prise), et 2 simples suggestions. Le 0 de Ministral est donc plus "
     "FIDÈLE que le 5 de Qwen. Autrement dit : extraire beaucoup peut vouloir dire "
     "halluciner beaucoup. La retenue d'un modèle est une qualité quand la réunion ne "
     "contient pas de décision.")
para("Côte à côte, même réunion, même passage (l'outil PPT de Maya, un projet PASSÉ "
     "qu'elle raconte). Qwen le transforme en décisions de la réunion ; Ministral "
     "indique correctement qu'aucune décision n'a été prise.")
side_by_side(
    "Qwen2.5 3B : fabrique des décisions",
    "« Décisions : Automatisation de la chaîne de traitement pour accélérer la "
    "production du rapport ; Définition d'une structure préétablie pour le rapport ». "
    "Or ce sont des éléments d'un projet ANTÉRIEUR raconté par Maya, pas des décisions "
    "prises en séance.",
    "archi_extraction_v4_qwen2.5-3b.md",
    "Ministral 3B : reste fidèle",
    "« ... Aucune décision prise concernant la validation ou l'extension de ces "
    "éléments. » (et en synthèse : « Si aucune décision claire n'a été prise pour "
    "orienter une mission spécifique... »)",
    "archi_extraction_v4_ministral3b.md",
    sentimentA="mauvais", sentimentB="bon",
)

h3("Modèles évalués (≤ 7B, CPU)")
table(
    ["Modèle", "Famille", "Quantization", "Taille"],
    [
        ["Ministral-3-3B-Instruct-2512", "Mistral 3B", "Q4_K_M", "~1,9 Go"],
        ["Mistral-7B-Instruct-v0.3", "Mistral 7B", "Q4_K_M", "~4,1 Go"],
        ["qwen2.5-3b-instruct", "Qwen 2.5 3B", "Q4_0", "~1,9 Go"],
        ["Qwen3-4B (thinking)", "Qwen 3 4B", "Q4_K_M", "~2,5 Go"],
        ["SmolLM3-3B", "HuggingFace", "Q4_K_M", "~1,9 Go"],
        ["LFM2.5-1.2B et LFM2-2.6B", "Liquid", "Q4 / Q6", "0,66 à 2,0 Go"],
        ["NuExtract-2.0-2B", "Extraction structurée", "Q4_K_M", "~1,3 Go"],
    ],
    widths=[2.4, 1.9, 1.2, 1.3],
)

h2("2.B  Refonte agentique : orchestrateur et workers (juin 2026)")
para("Limite de l'architecture extraction : elle impose la MÊME structure de CR à "
     "toutes les réunions. La refonte agentique vise une structure adaptée au type de "
     "réunion, via des agents spécialisés :", italic=True)
bullet("type de réunion (multi-hypothèses), objectif, synthèse globale.", "Agent CONTEXT BUILDER : ")
bullet("conçoit la structure (les sections).", "Agent PLANNER : ")
bullet("pour chaque section, choisit les chunks et rédige le brief.", "Agent CONTENT DESIGNER : ")
bullet("rédigent chaque section (narratif, puces, tableaux), à partir du texte brut.", "Agents WORKERS : ")
bullet("décisions et actions, puis assemblage Markdown 100 % Python.", "Juges déterministes : ")

h3("Évolution version par version (V1 à V10)")
para("Toutes sur la réunion RTE (dicte_audio_3). Une ligne est un run distinct ; la "
     "colonne résultat clé raconte ce que chaque version corrigeait.")
table(
    ["V", "Modèle(s)", "Sect.", "Durée totale", "Résultat clé, ce qu'elle corrigeait"],
    [
        ["V1", "Ministral 3B", "3", "20 min (1 225 s)", "1er jet : orchestrateur et workers"],
        ["V2", "Ministral 3B", "3", "26 min (1 569 s)", "archi multi-agents ; bug texte brut"],
        ["V3", "Ministral 3B", "4", "76 min (4 536 s)", "texte brut OK ; tour de table mal placé"],
        ["V4", "Ministral 3B", "4", "69 min (4 134 s)", "tables décisions/actions ; speaker mapping"],
        ["V5", "Qwen7B et draft 0.5B", "10", "incomplet", "test spec decoding, run interrompu"],
        ["V6", "Qwen 7B", "8", "146 min (8 758 s)", "prose en hausse mais détail PPT perdu"],
        ["V7", "Ministral 3B", "5", "182 min (10 940 s)", "1 worker planté (JSON vide), d'où robustesse"],
        ["V8", "Ministral 3B", "6", "125 min (7 525 s)", "robustesse et juges few-shot ; extraction fraîche"],
        ["V9", "Hybride 7B+3B", "7", "112 min (6 701 s)", "routage par agent (contexte 7B, reste 3B)"],
        ["V10", "Qwen 3B (sans diar.)", "10", "134 min (8 058 s)", "tout-Qwen, transcript sans diarisation"],
    ],
    widths=[0.4, 1.7, 0.5, 1.5, 2.7],
)
para("CR complets : agentique_v1 à v10_*.md (dossier comptes_rendus_references).",
     italic=True, size=8.5, color=BLEU2)

h3("Comparaison côte à côte : Ministral 3B (V8) vs Qwen 7B (V6)")
para("Même réunion, même section (automatisation de la rédaction de rapports). Le 7B "
     "cadre mieux globalement, mais sur le détail concret le 3B retient davantage (ici "
     "l'outil Power, les fichiers PPT riches en graphiques) là où le 7B reste générique.")
side_by_side(
    "Ministral 3B (V8) : détail conservé",
    "« Lors du projet mené dans l'équipe équilibre offre et demande, l'objectif "
    "consistait à automatiser la rédaction de rapports métiers détaillant les études de "
    "rentabilité des moyens de production. L'équipe utilisait déjà un outil interne "
    "nommé Power pour analyser ces données, mais les résultats étaient stockés sous "
    "forme de fichiers PPT contenant principalement des graphiques et peu de texte... »",
    "agentique_v8_ministral3b.md",
    "Qwen 7B (V6) : générique, détail perdu",
    "« La réunion a débuté par une présentation des rôles et expertises autour de "
    "projets liés à l'IA générative... Ensuite, Maya SAHRAOUI a présenté ses méthodes "
    "pour automatiser la rédaction de rapports à partir d'études PowerPoint. »",
    "agentique_v6_qwen7b.md",
    sentimentA="bon", sentimentB="mauvais",
)

h3("Exemples de défauts (même réunion) : pourquoi ces versions n'ont pas été retenues")
para("V10 (tout-Qwen, transcript sans diarisation) : Qwen sur-extrait et FABRIQUE des "
     "décisions, et faute de diarisation attribue tout au même nom (Bruno). Le plan "
     "d'action recopie même un exemple du prompt (Alice).")
extrait_cr(
    "agentique_v10_qwen3b_sans-diarisation.md",
    [
        "## 1. Discussion ... gestion des congestions",
        "| 1 | on décide de faire des connaissances | On décide de faire des connaissances |",
        "| 2 | on décide de gagner en souveraineté  | On décide de gagner en souveraineté |",
        "## 10. Plan d'action",
        "| 1 | S'occuper de la doc, vendredi | Alice | vendredi |   (Alice vient du prompt)",
    ],
    intro="Décisions fabriquées (la réunion n'en contient pas) et fuite d'un exemple du prompt :",
)
para("V7 (Ministral 3B) : un worker est parti en dérive de génération puis a renvoyé un "
     "JSON vide, donc section entièrement perdue. C'est ce run qui a déclenché les "
     "garde-fous (timeout, retry, fallback déterministe).")
extrait_cr(
    "agentique_v7_ministral3b.md",
    ["## 5. Plan d'action",
     "_Section non rendue (Expecting value: line 1 column 1 (char 0))._"],
    intro="Échec de robustesse (section vide) :",
)

h3("V9 : routage hybride par agent")
para("Constat des benchmarks : le 7B est meilleur UNIQUEMENT sur le cadrage global "
     "(Context Builder) et la fluidité ; le 3B suffit (et est plus rapide) ailleurs. V9 "
     "met donc un modèle différent par phase.")
table(
    ["Phase", "Modèle", "Durée"],
    [
        ["Context Builder", "Qwen 7B", "837 s"],
        ["Planner", "Ministral 3B", "629 s"],
        ["Content Designers", "Ministral 3B", "3 147 s"],
        ["Workers", "Ministral 3B", "2 040 s"],
        ["Swaps de modèle (coût)", "n.c.", "8,1 s"],
        ["Total", "n.c.", "112 min (6 701 s)"],
    ],
    widths=[2.3, 1.9, 1.6],
)
verdict("Prometteur (travail en cours). Le routage par agent fonctionne techniquement "
        "(swap de serveur environ 8 s, négligeable) : qualité du 7B sur le cadrage et "
        "vitesse du 3B sur le reste.")

doc.add_page_break()

# ===========================================================================
# PARTIE 3 — OPTIMISATIONS
# ===========================================================================
h1("Partie 3.  Les optimisations (vitesse et qualité)")
para("À architecture figée (boundary et extraction JSON, Ministral 3B Q4_K_M, "
     "dicte_audio_3.normalized), série d'A/B tests sur les techniques d'accélération de "
     "l'inférence et de réduction de coût. Une sous-partie est une technique.", italic=True)

h2("3.1  Baseline de départ")
para("Configuration initiale (prompts verbeux, 1 appel LLM par chunk).")
table(["Métrique", "Valeur"],
      [["Durée totale", "64 min 44 (3 883,5 s)"],
       ["Appels LLM", "8"],
       ["Génération sections", "3 189,8 s"],
       ["Executive summary", "338,0 s"],
       ["Plan d'attaque (LLM)", "308,5 s"]],
      widths=[3.0, 3.0])
verdict("Point de départ. Tout l'objectif des expériences est de faire descendre ce temps sans perdre en qualité.")

h2("3.2  Plan d'action : legacy vs perchunk (qualité)")
para("legacy = 1 appel LLM final sur les résumés agrégés. perchunk = engagements et "
     "suggestions extraits chunk par chunk (2 appels par chunk, bénéficiant du cache) "
     "puis assemblés en Python (aucun appel final).")
table(
    ["Mode", "Durée", "Appels LLM", "Items de plan extraits"],
    [
        ["legacy", "35 min 47 (2 147 s)", "16 (2/chunk)", "4"],
        ["perchunk", "40 min 33 (2 433 s)", "29 (4/chunk)", "17 (9 engagements + 8 suggestions)"],
    ],
    widths=[1.2, 1.8, 1.6, 2.4],
    green={(1, 3)}, red={(0, 3)},
)
para("CR complets : opt_plan-action_legacy_ministral3b.md et opt_plan-action_perchunk_ministral3b.md",
     italic=True, size=8.5, color=BLEU2)
verdict("Adopté (perchunk). Plus lent d'environ 5 min mais extrait 17 items au lieu de "
        "4, et l'assemblage final ne coûte plus d'appel LLM. Répond au reproche pas d'action prise.")

h2("3.3  Flash Attention : ON vs OFF (vitesse)")
para("Flash Attention est une implémentation optimisée du mécanisme d'attention "
     "(prérequis aussi pour la quantification du cache KV).")
table(
    ["Configuration", "Durée", "RAM serveur pic", "Chunks", "Appels LLM"],
    [
        ["Flash Attention ON", "25 min 14 (1 513,8 s)", "4 557 Mo", "10", "31"],
        ["Flash Attention OFF", "33 min 29 (2 009,2 s)", "5 467 Mo", "7", "22"],
    ],
    widths=[1.9, 1.8, 1.4, 1.0, 1.1],
    green={(0, 1)}, red={(1, 1)},
)
note("Le découpage a donné 10 chunks (ON) vs 7 (OFF), donc la comparaison n'est pas "
     "strictement iso-charge. Malgré plus de chunks, la version ON est plus rapide. "
     "CR : opt_flash-attention_on/off_ministral3b.md.")
verdict("Adopté (Flash Attention ON). Plus rapide (environ -25 %) et moins de RAM pic.")
pourquoi("L'attention standard matérialise une matrice N×N (N = nombre de tokens) qu'il "
         "faut écrire puis relire en mémoire, coûteux en bande passante (le vrai goulot). "
         "FlashAttention (Dao et al., 2022) calcule l'attention par TUILES sans jamais "
         "matérialiser cette matrice complète (algorithme IO-aware), d'où moins d'accès "
         "mémoire et moins de RAM. Dans llama.cpp, activer Flash Attention fournit ce "
         "noyau d'attention fusionné ET débloque la quantification du cache KV (section "
         "suivante), d'où le double gain temps et RAM observé.")
sources([
    ("FlashAttention, Dao, Fu, Ermon, Rudra, Ré, NeurIPS 2022", "https://arxiv.org/abs/2205.14135"),
])

h2("3.4  Cache KV quantifié (q8_0)")
para("Le cache KV (Key/Value) stocke le contexte déjà calculé. Le quantifier en q8_0 "
     "(via --cache-type-k/v q8_0, rendu possible par Flash Attention) réduit la bande "
     "passante mémoire, facteur limitant sur CPU.")
verdict("Adopté dans le pipeline. Allège la mémoire du cache sans dégradation observée de la sortie.")
pourquoi("Pendant la génération, le cache KV grossit avec la longueur du contexte et "
         "finit par dominer la consommation mémoire ; or sur CPU le facteur limitant est "
         "la bande passante mémoire (il faut relire ce cache à chaque token). Le "
         "quantifier en 8 bits (q8_0) divise environ par deux son empreinte et la bande "
         "passante nécessaire, avec une perte de qualité négligeable à 8 bits, c'est le "
         "constat général de la littérature sur la quantification du cache KV (KVQuant).")
sources([
    ("KVQuant, Hooper et al., NeurIPS 2024", "https://arxiv.org/abs/2401.18079"),
])

h2("3.5  Ordre du prompt document-first (vitesse, cache KV)")
para("Idée : placer le TEXTE du chunk AVANT les instructions. Comme on enchaîne 2 ou 3 "
     "appels sur le même chunk (résumé, puis extraction, puis plan), mettre le document "
     "en tête fait que son cache KV est calculé UNE fois (au 1er appel) puis RÉUTILISÉ "
     "par les appels suivants, qui n'ont plus qu'à traiter le petit bloc d'instructions "
     "qui change.")
para("Preuve chiffrée sur le même chunk (chunk 9, run ik_v5, log llama-server) :", bold=True)
table(
    ["Appel sur le chunk", "Tokens à prefiller", "Temps de prefill", "Pourquoi"],
    [
        ["1er, Résumé", "1 394 tokens", "~45,5 s", "charge tout le texte du chunk"],
        ["2e, Extraction", "220 tokens", "~9,3 s", "texte déjà en cache, ne traite que les instructions"],
    ],
    widths=[1.9, 1.6, 1.4, 2.3],
    red={(0, 2)}, green={(1, 2)},
)
note("Soit environ 6 fois moins de tokens à recalculer au 2e appel. La génération token "
     "par token reste à environ 7,7 tok/s dans les deux cas (le gain porte sur le "
     "prefill, pas la génération). Qualité de sortie identique entre les deux ordres.")
verdict("Adopté comme convention (document-first). Gain de temps gratuit sur les appels "
        "répétés d'un même chunk, sans aucun coût qualité.")
pourquoi("L'attention est causale : la représentation (cache KV) d'un token ne dépend "
         "que des tokens qui le PRÉCÈDENT. En plaçant le document en tête, on en fait un "
         "PRÉFIXE stable, identique pour les appels résumé, extraction et plan du même "
         "chunk : son cache KV est calculé une fois puis réutilisé, et seul le suffixe "
         "d'instructions (qui change) doit être recalculé. C'est le principe de la "
         "réutilisation de cache par préfixe formalisé par Prompt Cache (Gim et al., "
         "MLSys 2024). À l'inverse, instructions en tête casserait le préfixe commun et "
         "forcerait à tout recalculer.")
sources([
    ("Prompt Cache, Gim et al., MLSys 2024", "https://arxiv.org/abs/2311.04934"),
])

h2("3.6  Backend d'inférence : ik_llama.cpp vs llama.cpp mainline (vitesse et RAM)")
para("ik_llama.cpp est un fork de llama.cpp optimisé CPU. Comparé au mainline avec Flash Attention.")
table(
    ["Backend", "Durée", "RAM serveur (moy / pic)", "Décisions/Actions"],
    [
        ["mainline et FA", "25 min 14 (1 513,8 s)", "4 250 / 4 557 Mo", "0 / 6"],
        ["ik_llama v5", "22 min 41 (1 362,5 s)", "2 597 / 3 128 Mo", "3 / 9"],
    ],
    widths=[1.6, 1.8, 2.0, 1.4],
    green={(1, 1), (1, 2)},
)
note("CR complet : opt_ik-llama_v5_ministral3b.md. L'écart d'items extraits tient surtout "
     "à la variabilité du modèle, pas au backend.")
verdict("Prometteur. ik_llama plus rapide (environ -10 %) et surtout nettement plus "
        "économe en RAM (-1,4 Go pic). Le mainline reste retenu en production pour sa compatibilité.")
pourquoi("ik_llama.cpp est un fork de llama.cpp (par I. Kawrakow, à l'origine d'une "
         "partie des quantifications de llama.cpp). Il apporte des types de "
         "quantification SOTA (familles IQ et K-quants améliorées, meilleure perplexité à "
         "nombre de bits égal), des noyaux CPU SIMD plus rapides et un row-interleaved "
         "packing des poids qui améliore l'efficacité du cache processeur, d'où des gains "
         "surtout visibles en CPU et sur le traitement du prompt. Pas de publication "
         "académique : c'est un projet open-source (on cite donc le dépôt).")
sources([
    ("ik_llama.cpp, dépôt GitHub (ikawrakow)", "https://github.com/ikawrakow/ik_llama.cpp"),
])

h2("3.7  Speculative decoding (test dédié)")
para("Un petit modèle draft propose des tokens, le gros modèle cible les valide par "
     "lots (sortie mathématiquement identique). Micro-bench sur 1 chunk (chunk 4, "
     "14 631 caractères de prompt).")
table(
    ["Configuration", "Décodage", "Temps mur", "Acceptance"],
    [
        ["A, Ministral 3B seul", "4,82 tok/s", "268 s", "n.c."],
        ["B, Qwen 7B seul", "3,75 tok/s", "426 s", "n.c."],
        ["C, Qwen 7B et draft Qwen 0,5B", "2,00 tok/s", "472 s", "0,64 (61/95)"],
    ],
    widths=[2.6, 1.4, 1.2, 1.4],
    green={(0, 1)}, red={(2, 1), (2, 2)},
)
verdict("Rejeté dans notre configuration (Qwen 7B cible et draft 0,5B sur CPU). La "
        "config C est la PLUS LENTE (2,00 tok/s vs 3,75 pour le 7B seul).")
pourquoi("Le speculative decoding (Leviathan et al., ICML 2023) accélère un décodage "
         "MEMORY-BOUND : sur GPU à lot=1, générer un token est limité par la bande "
         "passante (charger les poids), pas par le calcul ; il reste donc du calcul "
         "disponible pour VÉRIFIER en parallèle les K tokens proposés par le draft "
         "presque gratuitement, d'où le 2 à 3 fois. Dans notre cas ce levier ne paie "
         "pas : (1) la vérification par lots sur CPU est limitée par le CALCUL (peu "
         "d'unités parallèles), vérifier K tokens coûte environ K passes, ce n'est plus "
         "gratuit ; (2) le draft tourne sur LES MÊMES cœurs que la cible et entre en "
         "concurrence ; (3) le taux d'acceptation (0,64) ne suffit pas à amortir ce "
         "surcoût sur une cible 7B. Honnêteté : les retours externes sur CPU sont "
         "mitigés (certains rapportent un gain sur des CPU très affamés en bande "
         "passante) ; notre mesure, sur cette paire de modèles précise, donne un ralentissement.")
sources([
    ("Speculative Decoding, Leviathan, Kalman, Matias, ICML 2023", "https://arxiv.org/abs/2211.17192"),
])

h2("3.8  Autres techniques testées (rejetées ou non concluantes)")
h3("Prompt Lookup Decoding (PLD)")
table(
    ["Étape", "PLD ON", "PLD OFF", "Écart"],
    [
        ["Extraction et plan", "210,20 s", "209,79 s", "+0,4 s (ON plus lent)"],
        ["Résumé", "224,70 s", "214,02 s", "+10,7 s (ON plus lent)"],
    ],
    widths=[1.8, 1.5, 1.5, 2.0],
)
verdict("Rejeté. Aucun gain, la sortie (résumé abstractif, JSON court) ne recopie pas assez le prompt.")
pourquoi("Le PLD remplace le modèle draft par une simple recopie de n-grammes déjà "
         "présents dans le prompt : il n'accélère que si la sortie REPREND littéralement "
         "l'entrée (résumé extractif, QA sur document, édition de code, fort recouvrement "
         "de n-grammes). Or notre sortie est un résumé ABSTRACTIF (reformulé) et un JSON "
         "court : le recouvrement de n-grammes avec le texte source est faible, il n'y a "
         "presque rien à deviner par copie, donc pas de gain.")
sources([
    ("Prompt Lookup Decoding, dépôt GitHub (apoorvumang)", "https://github.com/apoorvumang/prompt-lookup-decoding"),
])

h3("Compression de prompt (LLMLingua-2)")
table(
    ["Ratio", "Caractères", "Durée totale", "Effet qualité"],
    [
        ["1.0 (aucune)", "8 000", "308,7 s", "référence, extraction vide"],
        ["0.5", "4 152", "174,7 s", "identique à la référence"],
        ["0.4", "3 351", "128,7 s", "1 décision apparaît (extrapolation ?)"],
    ],
    widths=[1.2, 1.2, 1.5, 2.9],
)
verdict("Écarté. Gain de vitesse net (jusqu'à environ 2,9 fois) mais aux ratios "
        "agressifs le contenu extrait diverge (items non vérifiés). Risque qualité non "
        "maîtrisé sans gold.")
pourquoi("LLMLingua-2 (Pan et al., ACL 2024) compresse le prompt en classant chaque "
         "token à garder ou à supprimer (modèle entraîné par distillation de données). "
         "Moins de tokens veut dire moins à prefiller, donc plus rapide (gain réel "
         "environ 2,9 fois confirmé dans le papier). Mais supprimer des tokens peut "
         "retirer le contexte qui ancrait une affirmation : à ratio agressif (0,4) une "
         "décision est apparue dans notre test, exactement le type d'extrapolation "
         "interdit par notre contrainte zéro hallucination. Sans référence gold pour "
         "arbitrer, le risque n'est pas maîtrisable.")
sources([
    ("LLMLingua-2, Pan et al., Findings of ACL 2024", "https://aclanthology.org/2024.findings-acl.57/"),
])

h3("Parallélisme des slots et runtime OpenVINO")
bullet("Parallélisme (--parallel 1/2/4) : pas de gain réel sur CPU (cœurs partagés). Pipeline conservé séquentiel.")
bullet("OpenVINO : non concluant (prérequis lourds, Q4_K_M partiellement supporté, aucun résultat persisté).")

doc.add_page_break()

# ===========================================================================
# PARTIE 4 — SYNTHESE
# ===========================================================================
h1("Partie 4.  Synthèse")

h2("4.1  Bilan technique par technique")
table(
    ["Technique", "Partie", "Verdict"],
    [
        ["Découpage par frontières (vs HDBSCAN)", "1.1", "ADOPTÉ, chronologique, -32 à -41 %"],
        ["Extraction JSON et assemblage déterministe", "2.A", "ADOPTÉ, fiabilité du format"],
        ["Plan d'action perchunk", "3.2", "ADOPTÉ, 17 items vs 4"],
        ["Flash Attention", "3.3", "ADOPTÉ, environ -25 %, moins de RAM"],
        ["Cache KV quantifié (q8_0)", "3.4", "ADOPTÉ, bande passante CPU"],
        ["Prompt document-first (cache KV)", "3.5", "ADOPTÉ, prefill 6 fois moindre au 2e appel"],
        ["Backend ik_llama.cpp", "3.6", "PROMETTEUR, plus rapide, -1,4 Go, non déployé"],
        ["Speculative decoding", "3.7", "REJETÉ sur CPU, ralentit"],
        ["Prompt Lookup Decoding", "3.8", "REJETÉ, aucun gain"],
        ["Compression LLMLingua", "3.8", "ÉCARTÉ, rapide mais risque qualité"],
        ["Parallélisme des slots", "3.8", "REJETÉ, pas de gain CPU"],
        ["Runtime OpenVINO", "3.8", "NON CONCLU, prérequis lourds"],
        ["Architecture agentique", "2.B", "EN COURS, V8/V9"],
        ["Routage hybride par modèle", "2.B (V9)", "EN COURS, 7B cadrage et 3B reste"],
    ],
    widths=[2.7, 1.0, 3.1],
)

h2("4.2  Choix de modèle : ce que montrent les benchmarks")
bullet("le plus rapide des 3B et ose extraire, mais sur-extrait, donc fabrique des "
       "items quand la réunion n'en contient pas (à encadrer par un juge).", "Qwen2.5 3B : ")
bullet("modèle français de référence du pipeline actuel, fiable et sobre sur la "
       "rédaction (sa retenue est un atout anti-hallucination), plus faible sur le "
       "cadrage global du type de réunion.", "Ministral 3B : ")
bullet("meilleur cadrage global et fluidité, mais environ 2 fois plus lent ET il perd "
       "du détail concret (cf. comparaison V8 vs V6), réservé à un appel unique (Context "
       "Builder).", "Qwen 7B : ")
bullet("très rapides mais comptes rendus trop courts et pauvres.", "Modèles < 1,5B (LFM2.5, SmolLM) : ")

h2("4.3  État actuel et pistes")
para("Pipeline de production : architecture boundary et extraction JSON et assemblage "
     "déterministe, Ministral 3B Q4_K_M, Flash Attention ON, cache KV q8_0, plan "
     "perchunk, document-first, mode séquentiel. Environ 25 min sur une réunion d'1 h, "
     "moins de 5 Go de RAM, 100 % local.")
para("Exploration en cours : refonte agentique (V8/V9) pour adapter la structure du CR "
     "au type de réunion, avec routage hybride (7B pour le cadrage, 3B pour le reste) et "
     "garde-fous de robustesse (timeout, retry, fallback déterministe, plus un cap sur le "
     "nombre de tableaux pour éviter les plans dégénérés).")
para("Pistes ouvertes : backend ik_llama (RAM), amélioration de l'ASR et de la "
     "diarisation en amont (cause racine de plusieurs erreurs de noms), et, si la "
     "contrainte 100 % local était levée, une variante API cloud (map-reduce : "
     "extraction locale et synthèse cloud) plus rapide et de meilleure qualité.")

# ===========================================================================
# Sauvegarde
# ===========================================================================
out = Path(__file__).parent / "Rapport_Partie_LLM.docx"
try:
    doc.save(out)
    print(f"Document genere : {out}")
except PermissionError:
    alt = Path(__file__).parent / "Rapport_Partie_LLM_NEW.docx"
    doc.save(alt)
    print(f"[!] {out.name} est ouvert/verrouille -> enregistre sous : {alt}")
    print("    Ferme le fichier original puis relance pour ecraser le nom canonique.")
