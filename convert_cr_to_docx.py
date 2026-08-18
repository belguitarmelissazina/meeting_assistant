"""Convertit des comptes rendus Markdown en .docx (Word / Google Docs).

Gere : titres (#, ##, ###), tableaux GFM (| ... |), puces (-, *), listes
numerotees, gras **x**, italique *x* / _x_, code `x`, liens [t](url),
lignes de separation ---.

Par defaut, convertit les 3 comptes rendus compares :
  - _bench_mistral/compte_rendu_mistral.md
  - _bench_orchestrator_v8_3b/compte_rendu_v4.md
  - _bench_orchestrator_v9_hybride/compte_rendu_v9.md

Usage :
  python convert_cr_to_docx.py                       # les 3 par defaut
  python convert_cr_to_docx.py chemin1.md chemin2.md # fichiers au choix
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BLEU = RGBColor(0x1F, 0x3A, 0x5F)
GRIS = RGBColor(0x55, 0x55, 0x55)

_MD_ESCAPE_RE = re.compile(r"\\([\\`*_{}\[\]()#+\-.!|>])")


def _unescape_md(text: str) -> str:
    return _MD_ESCAPE_RE.sub(r"\1", text)


def _is_table_row(s: str) -> bool:
    s = s.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_separator(s: str) -> bool:
    s = s.strip()
    if not _is_table_row(s):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_table_row(s: str) -> list[str]:
    inner = s.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    parts = re.split(r"(?<!\\)\|", inner)
    return [p.strip().replace(r"\|", "|") for p in parts]


def _add_inline_runs(paragraph, text: str) -> None:
    text = _unescape_md(text)
    pattern = re.compile(
        r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|_[^_\n]+_|`[^`\n]+`|\[[^\]]+\]\([^)]+\))"
    )
    for tok in pattern.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("_") and tok.endswith("_") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            paragraph.add_run(m.group(1) if m else tok)
        else:
            paragraph.add_run(tok)


def _render_table(doc, header: list[str], body: list[list[str]]) -> None:
    ncols = max(len(header), max((len(r) for r in body), default=0))
    if ncols == 0:
        return
    t = doc.add_table(rows=1 + len(body), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        try:
            t.style = "Table Grid"
        except KeyError:
            pass
    hdr = t.rows[0].cells
    for j in range(ncols):
        cell = hdr[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header[j] if j < len(header) else "")
        run.bold = True
        run.font.size = Pt(9.5)
    for r_idx, row in enumerate(body, start=1):
        cells = t.rows[r_idx].cells
        for j in range(ncols):
            cell = cells[j]
            cell.text = ""
            _add_inline_runs(cell.paragraphs[0], row[j] if j < len(row) else "")
            for rn in cell.paragraphs[0].runs:
                rn.font.size = Pt(9.5)
    doc.add_paragraph()


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        # Tableau GFM
        if (_is_table_row(stripped) and i + 1 < len(lines)
                and _is_table_separator(lines[i + 1])):
            header = [_unescape_md(c) for c in _split_table_row(stripped)]
            i += 2
            body: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                body.append([_unescape_md(c) for c in _split_table_row(lines[i])])
                i += 1
            _render_table(doc, header, body)
            continue

        # Ligne de separation
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            i += 1
            continue

        # Titres
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            txt = _unescape_md(m.group(2))
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(txt)
                r.bold = True
                r.font.size = Pt(20)
                r.font.color.rgb = BLEU
            else:
                p = doc.add_heading(level=min(level, 4))
                r = p.add_run(txt)
                r.font.color.rgb = BLEU if level == 2 else GRIS
                r.font.size = Pt(14 if level == 2 else 12)
            i += 1
            continue

        # Puces
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # Listes numerotees
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    try:
        doc.save(str(docx_path))
        print(f"[OK] {md_path.name}  ->  {docx_path}")
    except PermissionError:
        alt = docx_path.with_name(docx_path.stem + "_NEW.docx")
        doc.save(str(alt))
        print(f"[!] {docx_path.name} verrouille (ouvert ?) -> {alt}")


DEFAULTS = [
    "_bench_mistral/compte_rendu_mistral.md",
    "_bench_orchestrator_v8_3b/compte_rendu_v4.md",
    "_bench_orchestrator_v9_hybride/compte_rendu_v9.md",
]


def main() -> int:
    here = Path(__file__).parent
    args = sys.argv[1:]
    targets = [Path(a) for a in args] if args else [here / d for d in DEFAULTS]
    n_ok = 0
    for md in targets:
        if not md.exists():
            print(f"[SKIP] introuvable : {md}", file=sys.stderr)
            continue
        md_to_docx(md, md.with_suffix(".docx"))
        n_ok += 1
    print(f"\n{n_ok}/{len(targets)} compte(s) rendu(s) converti(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
