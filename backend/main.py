"""
Backend FastAPI — diarisation & compte rendu.

Run:
    uvicorn backend.main:app --reload --port 8000

Workflow :
    1. POST /api/record/stop ou /api/process/upload → sauvegarde audio en "draft"
    2. POST /api/jobs/{id}/process + body {participants, entreprises, contexte}
       → lance le pipeline (diarisation, transcription, compte rendu)

Stockage : ~/Documents/Réunions/YYYY-MM-DD_HHhMMmSS/
    - audio.<ext>, audio.normalized.txt, compte_rendu.md
    - conservation illimitée — à l'utilisateur de faire le ménage
"""
from __future__ import annotations

# BLAS determinism — doit être set AVANT tout import numpy/scipy/sklearn (NMESC)
import os as _os
for _v in ("OMP_NUM_THREADS", "MKL_NUM_THREADS",
           "OPENBLAS_NUM_THREADS", "NUMEXPR_NUM_THREADS"):
    _os.environ.setdefault(_v, "1")

import asyncio
import ctypes
import logging
import shutil
import sys
import time
import uuid
from pathlib import Path
from typing import Literal, TYPE_CHECKING

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from pydantic import BaseModel

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

# NB : `audio_capture` tire numpy + (transitivement) sounddevice / sherpa /
# torch. C'est inutile pour démarrer le serveur HTTP et répondre à
# /api/health — on diffère donc l'import au premier `record_start()`. Grâce à
# `from __future__ import annotations` (haut du module), les annotations
# `recorder: AudioRecorder | None` ne sont jamais évaluées au runtime.
if TYPE_CHECKING:
    from audio_capture import AudioRecorder, LiveProcessor  # noqa: F401

from .job_logger import JobLogger, get_active as _get_active_log, set_active as _set_active_log  # noqa: E402

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("backend")


# ── Toggles debug (logs job + monitoring RAM) ─────────────────────────────
# Build packagé (PyInstaller frozen) → DÉSACTIVÉ EN DUR, sans aucun moyen
# de réactiver. Aucune variable d'environnement n'est lue dans ce mode :
# les utilisateurs finaux ne peuvent ni écrire `traitement_*.md` ni
# `resource_*.csv` dans leur dossier `~/.meeting_assistant/logs/`.
#
# Dev (interpréteur Python normal) → ON par défaut, désactivable via env
# vars `MEETING_JOB_LOG=0` / `MEETING_RAM_MONITOR=0`.
def _env_bool(name: str, default: bool) -> bool:
    val = _os.environ.get(name, "").strip().lower()
    if not val:
        return default
    return val not in ("0", "false", "no", "off")


_FROZEN = getattr(sys, "frozen", False)
if _FROZEN:
    # Build distribué : OFF par défaut. Les utilisateurs finaux ne
    # produisent ni traitement_*.md ni resource_*.csv. MAIS on garde une
    # trappe `MEETING_DEV_LOGS=1` pour le développeur qui veut debug
    # une install packagée — env var non documentée pour l'utilisateur.
    _DEV_OVERRIDE = _env_bool("MEETING_DEV_LOGS", default=False)
    _LOGS_ENABLED = _DEV_OVERRIDE
    _RAM_MONITOR_ENABLED = _DEV_OVERRIDE
else:
    _LOGS_ENABLED = _env_bool("MEETING_JOB_LOG", default=True)
    _RAM_MONITOR_ENABLED = _env_bool("MEETING_RAM_MONITOR", default=True)
log.info("Debug toggles : job_log=%s, ram_monitor=%s (frozen=%s)",
         _LOGS_ENABLED, _RAM_MONITOR_ENABLED, _FROZEN)

# ── Persistent storage — Documents/Réunions/ ──────────────────────────────
def _resolve_documents() -> Path:
    """Retourne le vrai dossier Documents (suivant OneDrive KFM si actif).

    Utilise l'API Windows SHGetFolderPathW pour obtenir le chemin vu par
    l'Explorateur — transparent sur tout PC (OneDrive ou non).
    """
    if sys.platform == "win32":
        try:
            import ctypes
            from ctypes import wintypes
            CSIDL_PERSONAL = 5
            SHGFP_TYPE_CURRENT = 0
            buf = ctypes.create_unicode_buffer(wintypes.MAX_PATH)
            res = ctypes.windll.shell32.SHGetFolderPathW(
                None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf
            )
            if res == 0 and buf.value:
                return Path(buf.value)
        except Exception as e:
            log.warning("SHGetFolderPathW a échoué : %s", e)
    return Path.home() / "Documents"


HISTORY_DIR = _resolve_documents() / "Réunions"
HISTORY_DIR.mkdir(parents=True, exist_ok=True)
log.info("Dossier des réunions : %s", HISTORY_DIR.resolve())

FOLDER_FMT = "%Y-%m-%d_%Hh%Mm%Ss"

# Paramètres locaux (clé API Mistral, etc.) — rangés hors de Documents
# pour éviter la synchro OneDrive.
SETTINGS_DIR = Path.home() / ".meeting_assistant"
SETTINGS_FILE = SETTINGS_DIR / "settings.json"


def _load_settings() -> dict:
    if not SETTINGS_FILE.exists():
        return {}
    try:
        import json as _json
        return _json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        log.warning("Settings illisibles (%s) — on repart à zéro", exc)
        return {}


def _save_settings(data: dict) -> None:
    import json as _json
    SETTINGS_DIR.mkdir(parents=True, exist_ok=True)
    SETTINGS_FILE.write_text(_json.dumps(data, indent=2), encoding="utf-8")


def _force_rmtree(path: Path) -> list[str]:
    """Supprime un dossier récursivement en forçant le passage des fichiers
    en lecture seule (attribut Windows fréquent sur les .docx synchronisés).

    Renvoie la liste des chemins qui n'ont pas pu être supprimés (ex : fichier
    verrouillé par Word/OneDrive). Un dossier entièrement supprimé renvoie [].
    """
    failures: list[str] = []

    def _on_error(func, p, exc_info):
        # Retry après avoir réinitialisé les droits d'écriture.
        try:
            import stat as _stat
            _os.chmod(p, _stat.S_IWRITE | _stat.S_IREAD)
            func(p)
        except Exception as exc:
            failures.append(f"{p} ({exc.__class__.__name__}: {exc})")

    try:
        shutil.rmtree(path, onerror=_on_error)
    except Exception as exc:
        failures.append(f"{path} ({exc.__class__.__name__}: {exc})")
    return failures


# ── Sleep prevention (Windows) ─────────────────────────────────────────────
ES_CONTINUOUS      = 0x80000000
ES_SYSTEM_REQUIRED = 0x00000001


def _prevent_sleep() -> None:
    if sys.platform == "win32":
        try:
            ctypes.windll.kernel32.SetThreadExecutionState(
                ES_CONTINUOUS | ES_SYSTEM_REQUIRED
            )
        except Exception:
            pass


def _allow_sleep() -> None:
    if sys.platform == "win32":
        try:
            ctypes.windll.kernel32.SetThreadExecutionState(ES_CONTINUOUS)
        except Exception:
            pass


# ── In-memory state ────────────────────────────────────────────────────────
JobStatus = Literal["draft", "pending", "queued", "running", "done", "error"]


JobSource = Literal["audio", "transcript"]
JobOrigin = Literal["recording", "upload"]


class Job:
    def __init__(
        self,
        job_id: str,
        source_path: Path,
        folder: Path | None = None,
        source: JobSource = "audio",
        origin: JobOrigin = "upload",
    ) -> None:
        self.id = job_id
        self.status: JobStatus = "draft"
        self.step = "Prêt à traiter"
        self.source: JobSource = source
        # Audio enregistré DANS l'app → clustering bootstrap+online (mode "live").
        # Audio/transcript uploadé   → clustering batch NMESC classique.
        self.origin: JobOrigin = origin
        # For `audio` source, `audio_path` is the audio file itself.
        # For `transcript` source, it's the uploaded raw transcript (.txt).
        self.audio_path = source_path
        self.out_dir = folder or source_path.parent
        self.out_dir.mkdir(parents=True, exist_ok=True)
        self.report_markdown: str | None = None
        self.transcript: str | None = None
        self.report_path: Path | None = None
        self.report_docx_path: Path | None = None
        self.transcript_path: Path | None = None
        self.error: str | None = None
        self.context: dict = {}
        # Réunion d'agenda liée (None si enregistrement hors agenda). Persisté
        # dans `.calendar_event.json` au sein du dossier → rechargé au boot.
        self.calendar: dict | None = None
        # Dossier (catégorie) classant la réunion, = sous-dossier réel de
        # Documents/Réunions/. None = à la racine ("Sans dossier").
        self.folder: str | None = None
        # st_ctime = creation time sur Windows, stable même après renommage
        self.created_at = self.out_dir.stat().st_ctime
        self.label = self.out_dir.name

    def public(self) -> dict:
        return {
            "id": self.id,
            "status": self.status,
            "step": self.step,
            "label": self.label,
            "source": self.source,
            "createdAt": int(self.created_at * 1000),
            "origin": self.origin,
            "audioAvailable": self.source == "audio" and self.audio_path.exists(),
            "reportMarkdown": self.report_markdown,
            "reportDocxAvailable": self.report_docx_path is not None and self.report_docx_path.exists(),
            "transcriptAvailable": self.transcript_path is not None and self.transcript_path.exists(),
            "transcript": self.transcript,
            "error": self.error,
            "context": self.context or {},
            "calendar": self.calendar,
            "folder": self.folder,
        }


jobs: dict[str, Job] = {}
recorder: AudioRecorder | None = None
live_processor: LiveProcessor | None = None
# Réunion d'agenda choisie au record/start, consommée au record/stop pour
# nommer le dossier + écrire le marqueur. None = enregistrement hors agenda.
_pending_calendar: dict | None = None
# Timestamp (ms epoch) du record/start en cours — utilisé par
# GET /api/record/status pour que le frontend reconstitue le timer côté UI
# même quand il rouvre la page après une navigation.
_recording_started_at_ms: int | None = None
pipeline_lock = asyncio.Lock()


def _looks_like_meeting(folder: Path) -> bool:
    """Un dossier de réunion contient un audio.* ou un transcript.raw.txt.
    Un dossier-catégorie n'en a pas (il ne contient que des sous-dossiers)."""
    try:
        if any(p.stem == "audio" and p.is_file() for p in folder.iterdir()):
            return True
    except OSError:
        return False
    return (folder / "transcript.raw.txt").exists()


def list_category_folders() -> list[str]:
    """Sous-dossiers de Documents/Réunions/ qui servent de catégories
    (= ne sont pas eux-mêmes une réunion). Inclut les dossiers vides."""
    out: list[str] = []
    try:
        for d in sorted(HISTORY_DIR.iterdir()):
            if d.is_dir() and not _looks_like_meeting(d):
                out.append(d.name)
    except OSError:
        pass
    return out


def _load_job_from_dir(folder: Path, category: str | None) -> tuple[str, Job] | None:
    audios = [p for p in folder.iterdir() if p.stem == "audio" and p.is_file()]
    raw_transcript = folder / "transcript.raw.txt"
    # Marqueur invisible : si le dossier a été créé par un enregistrement
    # live, on a posé ce fichier pour s'en souvenir au redémarrage.
    origin_marker = folder / ".origin.recording"
    origin: JobOrigin = "recording" if origin_marker.exists() else "upload"
    if audios:
        job_id = uuid.uuid4().hex
        job = Job(job_id, audios[0], folder, source="audio", origin=origin)
    elif raw_transcript.exists():
        job_id = uuid.uuid4().hex
        # Un transcript uploadé est toujours "upload" (pas de live pour du texte).
        job = Job(job_id, raw_transcript, folder, source="transcript", origin="upload")
    else:
        return None
    job.folder = category
    cal_marker = folder / ".calendar_event.json"
    if cal_marker.exists():
        try:
            import json as _json
            job.calendar = _json.loads(cal_marker.read_text(encoding="utf-8"))
        except (OSError, ValueError):
            job.calendar = None
    report_md = folder / "compte_rendu.md"
    report_docx = folder / "compte_rendu.docx"
    if report_md.exists() or report_docx.exists():
        job.status = "done"
        job.step = "Terminé"
        if report_md.exists():
            job.report_path = report_md
            try:
                job.report_markdown = report_md.read_text(encoding="utf-8")
            except OSError:
                pass
        if report_docx.exists():
            job.report_docx_path = report_docx
        if job.source == "audio":
            # Nouveau nom d'abord, puis ancien (dossiers créés avant le renommage).
            for candidate in (folder / "transcript.txt",
                              folder / "audio.transcript.midpoint.txt"):
                if candidate.exists():
                    job.transcript_path = candidate
                    try:
                        job.transcript = candidate.read_text(encoding="utf-8")
                    except OSError:
                        pass
                    break
            # Balayage unique au boot : épure les anciens dossiers encombrés.
            # UNIQUEMENT source="audio" — un dossier source="transcript" repose
            # sur transcript.raw.txt (cf. _looks_like_meeting) : ne pas l'effacer.
            try:
                _cleanup_intermediate_files(job)
            except Exception:
                log.exception("Nettoyage rehydrate échoué pour %s", folder)
    return job_id, job


def _rehydrate_jobs() -> None:
    for entry in sorted(HISTORY_DIR.iterdir()):
        if not entry.is_dir():
            continue
        if _looks_like_meeting(entry):
            loaded = _load_job_from_dir(entry, None)
            if loaded:
                jobs[loaded[0]] = loaded[1]
            continue
        # Sinon : dossier-catégorie → on charge les réunions qu'il contient
        # (un seul niveau de profondeur — pas de sous-sous-dossiers).
        for sub in sorted(entry.iterdir()):
            if sub.is_dir() and _looks_like_meeting(sub):
                loaded = _load_job_from_dir(sub, entry.name)
                if loaded:
                    jobs[loaded[0]] = loaded[1]
    if jobs:
        log.info("Historique rechargé : %d réunion(s)", len(jobs))


# Rehydrate de l'historique : scan de ~/Documents/Réunions/ + lecture
# INTÉGRALE de chaque compte_rendu.md / transcript.txt. Coût proportionnel
# au nombre de réunions passées → on le déporte dans un thread daemon pour
# ne PAS bloquer l'import du module. uvicorn binde le port et /api/health
# répond immédiatement ; les jobs apparaissent dans l'UI dès le scan fini
# (le frontend poll /api/jobs en continu).
def _rehydrate_jobs_bg() -> None:
    try:
        _rehydrate_jobs()
    except Exception:
        log.exception("Rehydrate de l'historique a échoué")


import threading as _threading_boot  # noqa: E402

_threading_boot.Thread(
    target=_rehydrate_jobs_bg, daemon=True, name="rehydrate-jobs"
).start()


# ── FastAPI app ────────────────────────────────────────────────────────────
app = FastAPI(title="Diarisation & Compte rendu")

app.add_middleware(
    CORSMiddleware,
    # Electron loads the bundled frontend from file:// (origin "null") or
    # app://, and the dev server uses :3000.
    allow_origins=[
        "http://localhost:3000", "http://127.0.0.1:3000",
        "null",
        "app://-", "app://./",
    ],
    allow_origin_regex=r"^(file://|app://).*",
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def _log_errors_only(request, call_next):
    """Ne logge QUE les requêtes HTTP qui échouent (status >= 400), pour
    garder la console lisible. uvicorn access_log est désactivé dans
    run_app._run_server()."""
    response = await call_next(request)
    if response.status_code >= 400:
        log.warning("%s %s → %d",
                    request.method, request.url.path, response.status_code)
    return response


# ── Monitoring RAM (CSV toutes les 30s dans ~/.meeting_assistant/logs/) ────
from .resource_monitor import get_monitor  # noqa: E402


@app.on_event("startup")
async def _start_resource_monitor() -> None:
    if _RAM_MONITOR_ENABLED:
        get_monitor().start()
    else:
        log.info("[RES] Monitoring RAM désactivé (MEETING_RAM_MONITOR=0)")


@app.on_event("shutdown")
async def _stop_resource_monitor() -> None:
    if _RAM_MONITOR_ENABLED:
        get_monitor().stop()


class MeetingContext(BaseModel):
    participants: str | None = None
    entreprises: str | None = None
    contexte: str | None = None
    diarize: bool = True
    llm: Literal["local", "mistral"] = "local"


class CalendarMeta(BaseModel):
    """Snapshot d'un événement Graph, figé au moment de l'enregistrement.
    Persisté tel quel dans `.calendar_event.json`."""
    eventId: str | None = None
    subject: str | None = None
    start: str | None = None
    end: str | None = None
    location: str | None = None
    organizer: str | None = None
    attendees: list[str] = []


class RecordStartPayload(BaseModel):
    """Payload optionnel pour démarrer un enregistrement.

    enableLiveLlm : si True, démarre aussi le LLM live pendant la captation
      (charge MiniLM + llama-server en avance, ~20-30s de cold start).
      Défaut False : on a juste la transcription live + la diarisation, le
      LLM tournera au "Lancer le traitement" comme avant.

    participants/entreprises/contexte : seulement utilisés si enableLiveLlm
      est True, pour ancrer les entités dans les prompts du LLM live.

    calendar : si fourni, l'enregistrement est rattaché à cette réunion
      d'agenda (dossier nommé d'après le sujet + marqueur `.calendar_event.json`).
    """
    enableLiveLlm: bool = False
    participants: str | None = None
    entreprises: str | None = None
    contexte: str | None = None
    calendar: CalendarMeta | None = None


class SettingsPayload(BaseModel):
    mistralApiKey: str | None = None
    # Comportement « fermer = quitter » (True) ou « fermer = cacher dans le
    # tray » (False — défaut). Tray = notifications reçues même quand la
    # fenêtre est fermée, mais l'app continue de consommer ~350 Mo RAM.
    quitOnClose: bool | None = None
    # Lance automatiquement l'app au démarrage Windows (Login Items API).
    # Défaut False — c'est à l'utilisateur de choisir cette friction.
    launchAtStartup: bool | None = None


class RenamePayload(BaseModel):
    label: str


class ReportPayload(BaseModel):
    markdown: str


class FolderPayload(BaseModel):
    name: str


class MoveFolderPayload(BaseModel):
    # None / "" => racine de Documents/Réunions (catégorie "Sans dossier").
    folder: str | None = None


# ── Pipeline runner ────────────────────────────────────────────────────────
async def _run_subprocess(cmd: list[str], cwd: Path, job: Job, extra_env: dict | None = None) -> None:
    env = {**_os.environ, "PYTHONIOENCODING": "utf-8"}
    if extra_env:
        env.update({k: v for k, v in extra_env.items() if v is not None})

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        cwd=str(cwd),
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.STDOUT,
        env=env,
    )
    assert proc.stdout is not None
    active_log = _get_active_log()
    async for raw in proc.stdout:
        line = raw.decode("utf-8", errors="replace").rstrip()
        if line:
            tagged = f"[job {job.id[:8]}] {line}"
            try:
                print(tagged, flush=True)
            except UnicodeEncodeError:
                # Filet de sécurité si stdout est resté en cp1252 (ex.
                # console Windows old-school) : on encode/replace avant.
                safe = tagged.encode("utf-8", errors="replace").decode("ascii", errors="replace")
                print(safe, flush=True)
            # Tee dans le fichier traitement.md du job en cours.
            if active_log is not None and not active_log.closed:
                active_log.append_raw(tagged)

    code = await proc.wait()
    if code != 0:
        raise RuntimeError(f"{cmd[0]} exited with code {code}")


async def run_pipeline(job: Job) -> None:
    try:
        job.status = "queued"
        job.step = "En file d'attente"
        async with pipeline_lock:
            job.status = "running"
            _prevent_sleep()
            try:
                await _run_pipeline_locked(job)
            finally:
                _allow_sleep()
    except Exception as e:
        log.exception("Pipeline job %s failed", job.id)
        job.status = "error"
        job.error = str(e)


def _subcmd(name: str, *args: str) -> list[str]:
    """Build a subcommand invocation for the multi-mode launcher.

    Frozen (PyInstaller) → reuse the same exe with a subcommand argument.
    Dev (regular interpreter) → call `python -m backend.run_app` with the
    same subcommand. Single dispatcher (`backend/run_app.py`), one code path.
    """
    if getattr(sys, "frozen", False):
        return [sys.executable, name, *args]
    python_bin = _os.environ.get("PYTHON_BIN") or sys.executable
    return [python_bin, "-u", "-m", "backend.run_app", name, *args]


async def _run_pipeline_locked(job: Job) -> None:
    # JobLogger optionnel (MEETING_JOB_LOG=0 → no-op, job_log = None).
    job_log: JobLogger | None = None
    if _LOGS_ENABLED:
        from .resource_monitor import LOG_DIR as _LOG_DIR
        _LOG_DIR.mkdir(parents=True, exist_ok=True)
        target_log_path = _LOG_DIR / f"traitement_{job.label}.md"
        job_log = _get_active_log()
        if job_log is None or job_log.closed or job_log.path != target_log_path:
            job_log = JobLogger(target_log_path, job.id, job.label)
            _set_active_log(job_log)
        else:
            # Réutilise le logger en cours (cas live → batch enchainés).
            job_log.job_id = job.id

    # Petits helpers pour traiter `job_log = None` (mode désactivé) sans
    # alourdir chaque call site avec un `if job_log is not None:`.
    def _log_start(name: str) -> None:
        if job_log is not None:
            job_log.start_step(name)

    def _log_end() -> None:
        if job_log is not None:
            job_log.end_step()

    try:
        if job.source == "audio":
            stem = job.audio_path.stem  # "audio"
            # Diarisation toujours active (plus d'option de désactivation côté
            # UI ; un client obsolète envoyant diarize=false est ignoré).
            diarize = True
            raw_transcript = job.out_dir / "transcript.txt"

            # Court-circuit live : si le LiveProcessor a déjà produit le
            # transcript pendant la captation, on saute toute l'étape diar.
            if raw_transcript.exists() and raw_transcript.stat().st_size > 0:
                job.step = "Transcript live prêt"
                log.info("Job %s : court-circuit diar (transcript live détecté)",
                         job.id)
            else:
                job.step = "Conversion + Diarisation + Transcription"
                _log_start(job.step)
                diar_args = ["-i", str(job.audio_path), "-o", str(job.out_dir)]
                if job.origin == "recording":
                    # Enregistrement dans l'app → clustering bootstrap+online.
                    diar_args.append("--bootstrap-online")
                await _run_subprocess(
                    _subcmd("diar", *diar_args),
                    PROJECT_ROOT, job,
                )
                _log_end()
                produced = job.out_dir / f"{stem}.transcript.midpoint.txt"
                if not produced.exists():
                    raise FileNotFoundError(f"Transcript introuvable : {produced}")
                # Renommage en un nom lisible dans l'Explorateur.
                if raw_transcript.exists():
                    raw_transcript.unlink()
                produced.rename(raw_transcript)
                # Le pipeline batch produit aussi `{stem}.turns.json` (turns
                # par locuteur avec start/end) — on le renomme en `turns.json`
                # pour que la vue Transcript le trouve (comme le fait le
                # LiveProcessor pour les enregistrements). Sans ça, l'audio
                # uploadé n'aurait pas d'onglet Transcript synchro audio.
                produced_turns = job.out_dir / f"{stem}.turns.json"
                if produced_turns.exists():
                    target_turns = job.out_dir / "turns.json"
                    if target_turns.exists():
                        target_turns.unlink()
                    produced_turns.rename(target_turns)
        else:
            # Transcript uploadé : on saute diar/transcription et on part
            # directement du fichier déjà converti en format standard.
            raw_transcript = job.audio_path

        job.step = "Normalisation transcript"
        _log_start(job.step)
        normalized = job.out_dir / "transcript.normalized.txt"
        await _run_subprocess(
            _subcmd("normalize", str(raw_transcript), str(normalized)),
            PROJECT_ROOT, job,
        )
        _log_end()
        # Pour les jobs audio : on expose le transcript issu de la transcription
        # (avec timestamps, tours de parole) pour téléchargement. Pour les jobs
        # source=transcript, l'utilisateur possède déjà son fichier.
        if job.source == "audio":
            job.transcript_path = raw_transcript
            try:
                job.transcript = raw_transcript.read_text(encoding="utf-8")
            except OSError:
                job.transcript = None

        job.step = "Compte rendu"
        report = job.out_dir / "compte_rendu.md"
        extra: list[str] = []
        ctx = job.context or {}
        if ctx.get("participants"):
            extra += ["--participants", ctx["participants"]]
        if ctx.get("entreprises"):
            extra += ["--entreprises", ctx["entreprises"]]

        llm_mode = ctx.get("llm", "local")
        subcmd_name = "mistral-minutes" if llm_mode == "mistral" else "minutes"
        extra_env = {
            "MEETING_CONTEXT":      ctx.get("contexte")     or "",
            "MEETING_PARTICIPANTS": ctx.get("participants") or "",
            "MEETING_ENTREPRISES":  ctx.get("entreprises")  or "",
        }
        if llm_mode == "mistral":
            api_key = (_load_settings().get("mistral_api_key") or "").strip()
            if not api_key:
                raise RuntimeError(
                    "Clé API Mistral absente. Renseigne-la dans les paramètres de l'application."
                )
            extra_env["MISTRAL_API_KEY"] = api_key
            job.step = "Compte rendu (Mistral Large)"

        _log_start(job.step)
        await _run_subprocess(
            _subcmd(subcmd_name,
                    "--transcript", str(normalized),
                    "--output", str(report),
                    *extra),
            PROJECT_ROOT, job,
            extra_env=extra_env,
        )
        _log_end()

        if report.exists():
            job.report_path = report
            job.report_markdown = report.read_text(encoding="utf-8")
            docx_path = job.out_dir / "compte_rendu.docx"
            try:
                _log_start("Conversion DOCX")
                _md_to_docx(report, docx_path)
                job.report_docx_path = docx_path
                _log_end()
            except Exception as exc:
                log.warning("Conversion DOCX échouée pour job %s : %s", job.id, exc)
                _log_end()

        _cleanup_intermediate_files(job)

        job.status = "done"
        job.step = "Terminé"
    finally:
        # Toujours fermer le JobLogger (même si exception) pour écrire le
        # tableau récap dans traitement.md et libérer le FileHandler.
        if job_log is not None and not job_log.closed:
            job_log.close()
            _set_active_log(None)


def _cleanup_intermediate_files(job: Job) -> None:
    """Ne garde dans le dossier réunion QUE : l'audio, compte_rendu.md/.docx,
    `.calendar_event.json` (lien agenda), et les fichiers de transcript par
    tours (transcript.txt + turns.json) qui servent l'UI « Transcript synchro
    audio + rename speakers ». Tout le reste (transcript normalisé pour le
    LLM, words.json, *.json diarisation, métriques…) est supprimé. Le
    `traitement_<label>.md` est ailleurs (~/.meeting_assistant/logs/)."""
    keep = {job.audio_path.name, "compte_rendu.md", "compte_rendu.docx",
            ".calendar_event.json",
            "transcript.txt", "turns.json", "speakers.json"}
    for entry in job.out_dir.iterdir():
        if not entry.is_file() or entry.name in keep:
            continue
        try:
            entry.unlink()
        except OSError as exc:
            log.warning("Suppression échouée pour %s : %s", entry, exc)


# ── Helpers ────────────────────────────────────────────────────────────────
def _new_folder(name_hint: str | None = None) -> Path:
    """Dossier daté unique ; ajoute un suffixe en cas de collision.

    `name_hint` (ex. sujet de la réunion d'agenda) → suffixe lisible dans
    l'Explorateur : `2026-05-19_14h00m00s_Revue produit`. Hors agenda, on
    garde l'horodatage seul comme avant.
    """
    base = time.strftime(FOLDER_FMT)
    if name_hint:
        safe = _sanitize_label(name_hint)
        if safe:
            base = f"{base}_{safe}"
    folder = HISTORY_DIR / base
    i = 2
    while folder.exists():
        folder = HISTORY_DIR / f"{base}_{i}"
        i += 1
    folder.mkdir(parents=True, exist_ok=True)
    return folder


_FORBIDDEN_LABEL_CHARS = set('<>:"/\\|?*')
_WINDOWS_RESERVED = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


def _sanitize_label(raw: str) -> str:
    cleaned = "".join(
        c for c in (raw or "").strip()
        if c not in _FORBIDDEN_LABEL_CHARS and ord(c) >= 32
    )
    cleaned = cleaned.strip(". ")
    if cleaned.upper() in _WINDOWS_RESERVED:
        cleaned = f"_{cleaned}"
    return cleaned[:200]


# ── Markdown -> DOCX ───────────────────────────────────────────────────────
_MD_ESCAPE_RE = __import__("re").compile(r"\\([\\`*_{}\[\]()#+\-.!|>])")


def _unescape_md(text: str) -> str:
    """Annule les échappements markdown (ex: `1\\.` -> `1.`, `_italic_`)."""
    return _MD_ESCAPE_RE.sub(r"\1", text)


def _is_table_row(s: str) -> bool:
    s = s.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_separator(s: str) -> bool:
    import re
    s = s.strip()
    if not _is_table_row(s):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_table_row(s: str) -> list[str]:
    import re
    inner = s.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    parts = re.split(r"(?<!\\)\|", inner)
    return [p.strip().replace(r"\|", "|") for p in parts]


def _md_to_docx(md_path: Path, docx_path: Path) -> None:
    import re
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Table GFM : ligne en pipes suivie d'un séparateur `|---|---|`
        if (
            _is_table_row(stripped)
            and i + 1 < len(lines)
            and _is_table_separator(lines[i + 1])
        ):
            header = [_unescape_md(c) for c in _split_table_row(stripped)]
            i += 2
            body: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                body.append([_unescape_md(c) for c in _split_table_row(lines[i])])
                i += 1
            _render_table(doc, header, body)
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_unescape_md(m.group(2)), level=level)
            i += 1
            continue

        if stripped.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\\?\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\\?\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))


def _render_table(doc, header: list[str], body: list[list[str]]) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL

    ncols = max(len(header), max((len(r) for r in body), default=0))
    if ncols == 0:
        return
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

    header_cells = table.rows[0].cells
    for j in range(ncols):
        cell = header_cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.text = ""
        text = header[j] if j < len(header) else ""
        run = para.add_run(text)
        run.bold = True

    for r_idx, row in enumerate(body, start=1):
        row_cells = table.rows[r_idx].cells
        for j in range(ncols):
            cell = row_cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cell.paragraphs[0]
            para.text = ""
            text = row[j] if j < len(row) else ""
            _add_inline_runs(para, text)


def _add_inline_runs(paragraph, text: str) -> None:
    """Gère **gras**, *italique*, _italique_, `code` et [texte](url).
    Les échappements markdown (`1\\.`, `\\*`) sont neutralisés en amont.
    """
    import re
    text = _unescape_md(text)
    pattern = re.compile(
        r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|_[^_\n]+_|`[^`\n]+`|\[[^\]]+\]\([^)]+\))"
    )
    for tok in pattern.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("_") and tok.endswith("_") and len(tok) > 2:
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if m:
                paragraph.add_run(m.group(1)).italic = True
            else:
                paragraph.add_run(tok)
        else:
            paragraph.add_run(tok)


# ── Teams .docx transcript parser ──────────────────────────────────────────
_TS_RE = __import__("re").compile(r"^\d{1,2}(?::\d{2}){1,2}$")


def _normalize_teams_timestamp(raw: str) -> str | None:
    """'0:08' → '00:00:08' ; '1:03:43' → '01:03:43'. None si non conforme."""
    if not _TS_RE.match(raw):
        return None
    parts = raw.split(":")
    if len(parts) == 2:
        return f"00:{parts[0].zfill(2)}:{parts[1]}"
    return f"{parts[0].zfill(2)}:{parts[1]}:{parts[2]}"


def _parse_teams_docx_transcript(docx_path: Path) -> str:
    """Convertit un transcript Teams exporté en .docx vers le format standard
    `[HH:MM:SS] Speaker: texte` attendu par normalize_transcript.py.

    Structure par paragraphe : run gras = nom du locuteur, run normal suivant
    = timestamp (MM:SS ou H:MM:SS), puis plusieurs runs séparés par <w:br/>
    = phrases successives. Les paragraphes de méta (titre, "X a commencé la
    transcription", etc.) n'ont pas de timestamp valide et sont ignorés.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    try:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml").decode("utf-8")
    except (zipfile.BadZipFile, KeyError) as exc:
        raise ValueError("Fichier .docx invalide") from exc

    root = ET.fromstring(xml)
    body = root.find("w:body", NS)
    if body is None:
        raise ValueError("Document Word vide")

    lines: list[str] = []
    for p in body.findall("w:p", NS):
        # (is_bold, text, is_break) tuples — ordre préservé
        segments: list[tuple[bool | None, str, bool]] = []
        for r in p.findall("w:r", NS):
            rpr = r.find("w:rPr", NS)
            is_bold = rpr is not None and rpr.find("w:b", NS) is not None
            for child in r:
                tag = child.tag
                if tag == W + "t":
                    segments.append((is_bold, child.text or "", False))
                elif tag == W + "br":
                    segments.append((None, "", True))

        speaker: str | None = None
        timestamp: str | None = None
        sentences: list[str] = []
        current: list[str] = []
        for is_bold, text, is_break in segments:
            if is_break:
                if current:
                    joined = "".join(current).strip()
                    if joined:
                        sentences.append(joined)
                    current = []
                continue
            if speaker is None:
                if is_bold and text.strip():
                    speaker = text.strip()
                continue
            if timestamp is None:
                ts = _normalize_teams_timestamp(text.strip())
                if ts is not None:
                    timestamp = ts
                    continue
                # Pas un timestamp → paragraphe de méta, on abandonne
                break
            current.append(text)
        if current:
            joined = "".join(current).strip()
            if joined:
                sentences.append(joined)

        if not speaker or not timestamp or not sentences:
            continue

        text = " ".join(sentences).strip()
        if not text:
            continue
        lines.append(f"[{timestamp}] {speaker}: {text}")

    if not lines:
        raise ValueError("Aucun tour de parole reconnu dans le document")

    return "\n".join(lines) + "\n"


# ── Endpoints ──────────────────────────────────────────────────────────────
@app.get("/api/health")
async def health() -> dict:
    return {"ok": True, "jobs": len(jobs), "recording": recorder is not None and recorder.is_recording}


@app.get("/api/settings")
async def get_settings() -> dict:
    """Expose l'état des préférences (ne renvoie JAMAIS la clé API en clair)."""
    s = _load_settings()
    key = (s.get("mistral_api_key") or "").strip()
    return {
        "mistralKeySet": bool(key),
        # Défaut FALSE = comportement tray opt-out : fermer la fenêtre laisse
        # l'app vivante en arrière-plan pour les notifications.
        "quitOnClose": bool(s.get("quit_on_close", False)),
        # Défaut FALSE = pas de friction au démarrage Windows.
        "launchAtStartup": bool(s.get("launch_at_startup", False)),
    }


@app.put("/api/settings")
async def update_settings(body: SettingsPayload) -> dict:
    """Enregistre/efface chaque champ INDIVIDUELLEMENT (None = no-op).

    - `mistralApiKey = "..."` → enregistre la clé Mistral.
    - `mistralApiKey = ""`   → efface la clé Mistral.
    - `quitOnClose` / `launchAtStartup` : bool → écrit le flag, None ignore.
    """
    s = _load_settings()
    if body.mistralApiKey is not None:
        cleaned = body.mistralApiKey.strip()
        if cleaned:
            s["mistral_api_key"] = cleaned
        else:
            s.pop("mistral_api_key", None)
    if body.quitOnClose is not None:
        s["quit_on_close"] = bool(body.quitOnClose)
    if body.launchAtStartup is not None:
        s["launch_at_startup"] = bool(body.launchAtStartup)
    _save_settings(s)
    key = (s.get("mistral_api_key") or "").strip()
    return {
        "mistralKeySet": bool(key),
        "quitOnClose": bool(s.get("quit_on_close", False)),
        "launchAtStartup": bool(s.get("launch_at_startup", False)),
    }


# ── Calendrier Microsoft (Graph, device code flow délégué) ─────────────────
# `graph_calendar` tire msal/requests → import PARESSEUX dans chaque handler
# pour ne pas alourdir le démarrage du serveur (cf. /api/health rapide).
# Les fonctions du module sont synchrones (msal/requests) → asyncio.to_thread.
@app.get("/api/calendar/status")
async def calendar_status() -> dict:
    from . import graph_calendar as gc
    return await asyncio.to_thread(gc.status)


@app.post("/api/calendar/login")
async def calendar_login() -> dict:
    """Démarre le device code flow. Renvoie le code à saisir sur
    microsoft.com/devicelogin ; l'acquisition se poursuit en tâche de fond
    (le front poll /api/calendar/status jusqu'à `signed_in`)."""
    from . import graph_calendar as gc
    return await asyncio.to_thread(gc.start_login)


@app.post("/api/calendar/logout")
async def calendar_logout() -> dict:
    from . import graph_calendar as gc
    await asyncio.to_thread(gc.logout)
    return {"state": "signed_out"}


@app.get("/api/calendar/upcoming")
async def calendar_upcoming(days: int = 7) -> JSONResponse:
    from . import graph_calendar as gc
    try:
        meetings = await asyncio.to_thread(gc.list_upcoming, days)
    except gc.NotAuthenticated as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except Exception as exc:
        log.warning("Lecture du calendrier échouée : %s", exc)
        raise HTTPException(status_code=502, detail=f"Calendrier indisponible : {exc}") from exc
    return JSONResponse({"meetings": meetings})


@app.post("/api/record/cancel")
async def record_cancel() -> dict:
    """Reset en cas d'enregistrement fantôme (backend coincé après un stop
    raté ou une init bloquée)."""
    global recorder, live_processor, _pending_calendar, _recording_started_at_ms
    _pending_calendar = None
    _recording_started_at_ms = None
    cleaned = False
    if recorder is not None:
        try:
            await asyncio.to_thread(recorder.stop)
        except Exception:
            pass
        recorder = None
        cleaned = True
    if live_processor is not None:
        try:
            live_processor._running = False
        except Exception:
            pass
        live_processor = None
        cleaned = True
    # Ferme le JobLogger s'il existait (le tmp file restera dans %TEMP% mais
    # au moins le FileHandler est libéré et le récap écrit).
    job_log = _get_active_log()
    if job_log is not None and not job_log.closed:
        job_log.close()
        _set_active_log(None)
        cleaned = True
    return {"cleaned": cleaned}


@app.get("/api/record/status")
async def record_status() -> dict:
    """État du recorder global (un seul enregistrement à la fois). Permet au
    frontend de re-synchroniser son UI : si l'utilisateur quitte la page
    pendant un enregistrement et y revient, le bouton/timer doivent refléter
    la réalité (backend qui enregistre toujours)."""
    if recorder is None or not recorder.is_recording:
        return {"recording": False}
    return {
        "recording": True,
        "startedAt": _recording_started_at_ms,
        "calendar": _pending_calendar,
    }


@app.post("/api/record/start")
async def record_start(payload: RecordStartPayload | None = None) -> dict:
    global recorder, live_processor, _pending_calendar, _recording_started_at_ms
    # Import différé : 1er endroit où l'on construit réellement AudioRecorder /
    # LiveProcessor. Sort le coût d'import (numpy/sherpa/…) du démarrage serveur.
    from audio_capture import AudioRecorder, LiveProcessor

    # Réunion d'agenda liée à cet enregistrement (consommée au record/stop).
    _pending_calendar = (
        payload.calendar.model_dump()
        if payload is not None and payload.calendar is not None
        else None
    )
    # Auto-reset si un enregistrement précédent est coincé (init ratée,
    # stop perdu, backend rechargé à chaud, …). Mieux vaut récupérer que
    # bloquer l'utilisateur sur un 409.
    if recorder is not None:
        try:
            await asyncio.to_thread(recorder.stop)
        except Exception:
            pass
        recorder = None
    if live_processor is not None:
        try:
            live_processor._running = False
        except Exception:
            pass
        live_processor = None

    # Mode live ACTIVÉ par défaut pour les enregistrements dans l'app :
    # diarisation + LLM tournent en parallèle pendant la captation. LLM
    # désactivable via payload (ex : dev/test sans llama-server).
    enable_live_llm = True
    if payload is not None and not payload.enableLiveLlm:
        # Opt-out explicite : si False fourni, on désactive juste le LLM
        # (on garde la transcription/diarisation live qui sont légères).
        enable_live_llm = False

    live_processor = LiveProcessor(enable_live_llm=enable_live_llm)
    if enable_live_llm and payload is not None:
        live_processor.set_llm_context({
            "participants": payload.participants or "",
            "entreprises": payload.entreprises or "",
            "contexte": payload.contexte or "",
        })

    # JobLogger optionnel : créé seulement si MEETING_JOB_LOG=1 (défaut dev).
    # Désactivé dans l'exe par défaut → zéro overhead de fichier .md.
    if _LOGS_ENABLED:
        from .resource_monitor import LOG_DIR as _LOG_DIR
        _LOG_DIR.mkdir(parents=True, exist_ok=True)
        pending_log_path = _LOG_DIR / f"traitement_pending_{int(time.time())}.md"
        job_log = JobLogger(pending_log_path, job_id="recording-pending",
                             label="Enregistrement en cours")
        _set_active_log(job_log)
        job_log.start_step(
            "Captation audio + traitement live "
            f"(transcription + LLM={'ON' if enable_live_llm else 'OFF'} en parallèle)"
        )

    # Branche le recorder maintenant : l'audio commence à être capté et
    # poussé dans les queues du LiveProcessor. Les workers démarreront
    # dans ~15-20s quand les modèles auront fini de charger — entre-temps
    # les chunks s'empilent sans être perdus.
    on_chunk = live_processor.push
    recorder = AudioRecorder(on_chunk=on_chunk)
    recorder.start()

    # Démarrage de LiveProcessor dans un thread de fond. La response
    # HTTP revient immédiatement, le chargement des modèles se fait en
    # parallèle de la captation.
    def _background_start(lp):
        try:
            lp.start()
            if lp.error:
                log.warning("LiveProcessor init fini avec erreur : %s", lp.error)
            else:
                log.info("LiveProcessor prêt (workers en route)")
        except Exception as exc:
            log.exception("Background start LiveProcessor a levé : %s", exc)
    import threading as _t
    _t.Thread(target=_background_start, args=(live_processor,),
              daemon=True, name="lp-start").start()

    _recording_started_at_ms = int(time.time() * 1000)
    return {
        "pid": _os.getpid(),
        "startedAt": _recording_started_at_ms,
        "liveProcessing": True,
        "liveLlm": enable_live_llm,
    }


@app.post("/api/record/stop")
async def record_stop() -> dict:
    global recorder, live_processor, _pending_calendar, _recording_started_at_ms
    if recorder is None or not recorder.is_recording:
        raise HTTPException(status_code=400, detail="Aucun enregistrement en cours")
    # Snapshot + reset de la réunion liée (posée au record/start).
    cal = _pending_calendar
    _pending_calendar = None
    _recording_started_at_ms = None
    # Marque le moment du clic STOP — chronos 'attente compte rendu' démarre.
    # Affiché en évidence dans le summary du traitement.md.
    job_log = _get_active_log()
    if job_log is not None and not job_log.closed:
        job_log.mark_stop_clicked()
    try:
        audio_path = await asyncio.to_thread(recorder.stop)
    finally:
        recorder = None
    if not audio_path:
        if live_processor is not None:
            try:
                await asyncio.to_thread(live_processor.finalize, Path("."))
            except Exception:
                pass
            live_processor = None
        # Pas d'audio → pas de dossier final → on ferme le tmp logger.
        job_log = _get_active_log()
        if job_log is not None and not job_log.closed:
            job_log.close()
            _set_active_log(None)
        raise HTTPException(status_code=500, detail="Aucune donnée audio capturée")

    folder = _new_folder(cal.get("subject") if cal else None)
    final_path = folder / "audio.wav"
    shutil.move(str(audio_path), str(final_path))
    # Marqueur de persistance : pour savoir au rehydrate que ce dossier
    # provient d'un enregistrement live (→ clustering bootstrap+online).
    (folder / ".origin.recording").write_text("", encoding="utf-8")
    # Marqueur de rattachement à la réunion d'agenda (rechargé au boot).
    if cal:
        import json as _json
        (folder / ".calendar_event.json").write_text(
            _json.dumps(cal, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    # Renomme le JobLogger avec le nom du dossier de la réunion (toujours
    # dans ~/.meeting_assistant/logs/, à côté du CSV RAM). No-op si le
    # logging est désactivé (MEETING_JOB_LOG=0).
    job_log = _get_active_log()
    if job_log is not None and not job_log.closed:
        from .resource_monitor import LOG_DIR as _LOG_DIR
        job_log.end_step()  # ferme "Captation + traitement live"
        job_log.label = folder.name
        job_log.move(_LOG_DIR / f"traitement_{folder.name}.md")
        job_log.start_step("Finalisation live (drain chunker + LLM restant + assemblage)")

    # Finalise le pipeline live : si réussi, écrit transcript.txt dans le
    # dossier. Le pipeline batch détectera la présence du fichier et
    # court-circuitera l'étape diar.
    live_ok = False
    if live_processor is not None:
        try:
            live_ok = await asyncio.to_thread(live_processor.finalize, folder)
        except Exception:
            log.exception("LiveProcessor.finalize a levé")
        finally:
            live_processor = None
    if live_ok:
        log.info("Transcript live écrit — diar batch sera court-circuité")
    else:
        log.info("Pas de transcript live — diar batch tournera au lancement")
    if job_log is not None and not job_log.closed:
        job_log.end_step()  # ferme "Finalisation live"

    job_id = uuid.uuid4().hex
    job = Job(job_id, final_path, folder, source="audio", origin="recording")
    job.calendar = cal

    # Si le LLM live a abouti pendant la captation, compte_rendu.md est
    # déjà dans le dossier → on marque le job "done" direct, l'utilisateur
    # n'a pas à cliquer "Lancer le traitement".
    report_md = folder / "compte_rendu.md"
    if report_md.exists() and report_md.stat().st_size > 0:
        job.status = "done"
        job.step = "Terminé (live)"
        job.report_path = report_md
        try:
            job.report_markdown = report_md.read_text(encoding="utf-8")
        except OSError:
            pass
        docx_path = folder / "compte_rendu.docx"
        try:
            _md_to_docx(report_md, docx_path)
            job.report_docx_path = docx_path
        except Exception as exc:
            log.warning("Conversion DOCX live KO : %s", exc)
        transcript = folder / "transcript.txt"
        if transcript.exists():
            job.transcript_path = transcript
            try:
                job.transcript = transcript.read_text(encoding="utf-8")
            except OSError:
                pass
        log.info("Job %s : compte rendu live prêt — aucune action manuelle requise",
                 job.id)
        # Le chemin live-done ne passe jamais par _run_pipeline_locked (où le
        # cleanup est sinon déclenché) → on épure ici. transcript déjà chargé
        # en mémoire (job.transcript) donc la vue en session reste OK.
        try:
            _cleanup_intermediate_files(job)
        except Exception:
            log.exception("Nettoyage live-done échoué pour %s", folder)

    # Si le job est totalement fini (compte rendu live OK) → on ferme le
    # JobLogger maintenant (récap + tableau étapes écrits dans traitement.md).
    # Sinon on le laisse ouvert : il sera repris par job_process pour
    # logger la pipeline batch quand l'utilisateur cliquera "Lancer".
    if job.status == "done":
        job_log = _get_active_log()
        if job_log is not None and not job_log.closed:
            job_log.job_id = job.id
            job_log.close()
            _set_active_log(None)

    jobs[job_id] = job
    return {
        "jobId": job_id,
        "liveTranscriptReady": live_ok,
        "liveReportReady": job.status == "done",
    }


@app.post("/api/process/upload")
async def process_upload(audio: UploadFile = File(...)) -> dict:
    ext = Path(audio.filename or "audio.wav").suffix or ".wav"
    folder = _new_folder()
    dest = folder / f"audio{ext}"
    with dest.open("wb") as f:
        shutil.copyfileobj(audio.file, f)

    job_id = uuid.uuid4().hex
    job = Job(job_id, dest, folder, source="audio", origin="upload")
    jobs[job_id] = job
    return {"jobId": job_id}


@app.post("/api/process/upload-transcript")
async def process_upload_transcript(transcript: UploadFile = File(...)) -> dict:
    """Upload d'un transcript déjà produit (Teams .docx ou .txt formaté).

    La chaîne normalisation → compte rendu est exécutée sans l'étape audio.
    """
    filename = transcript.filename or "transcript.txt"
    ext = Path(filename).suffix.lower()
    if ext not in {".txt", ".docx"}:
        raise HTTPException(
            status_code=400,
            detail="Format non supporté. Utilisez .txt ou .docx (Teams).",
        )

    folder = _new_folder()
    dest = folder / "transcript.raw.txt"
    try:
        if ext == ".docx":
            tmp_docx = folder / "_source.docx"
            with tmp_docx.open("wb") as f:
                shutil.copyfileobj(transcript.file, f)
            try:
                content = _parse_teams_docx_transcript(tmp_docx)
            finally:
                tmp_docx.unlink(missing_ok=True)
            dest.write_text(content, encoding="utf-8")
        else:
            data = transcript.file.read()
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError:
                text = data.decode("utf-8", errors="replace")
            if not text.strip():
                raise ValueError("Transcript vide")
            dest.write_text(text, encoding="utf-8")
    except ValueError as exc:
        shutil.rmtree(folder, ignore_errors=True)
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        shutil.rmtree(folder, ignore_errors=True)
        log.exception("Import du transcript échoué")
        raise HTTPException(status_code=500, detail=f"Import impossible : {exc}") from exc

    job_id = uuid.uuid4().hex
    job = Job(job_id, dest, folder, source="transcript")
    jobs[job_id] = job
    return {"jobId": job_id}


@app.post("/api/jobs/{job_id}/process")
async def job_process(
    job_id: str,
    background: BackgroundTasks,
    context: MeetingContext,
) -> dict:
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.status in ("queued", "running"):
        raise HTTPException(status_code=409, detail="Traitement déjà en cours")
    if not job.audio_path.exists():
        missing = "Audio" if job.source == "audio" else "Transcript"
        raise HTTPException(status_code=404, detail=f"{missing} introuvable")

    # Idempotence : si le compte rendu est déjà présent (écrit par le LLM
    # live), pas besoin de relancer — on renvoie simplement le status.
    if job.status == "done" and job.report_path and job.report_path.exists():
        job.context = context.model_dump()
        return {"jobId": job.id, "status": job.status, "alreadyDone": True}

    job.context = context.model_dump()
    job.status = "pending"
    job.step = "Initialisation"
    job.error = None
    background.add_task(run_pipeline, job)
    return {"jobId": job.id, "status": job.status}


@app.get("/api/jobs")
async def list_jobs() -> JSONResponse:
    items = sorted(jobs.values(), key=lambda j: j.created_at, reverse=True)
    return JSONResponse({"jobs": [j.public() for j in items]})


@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str) -> JSONResponse:
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    return JSONResponse(job.public())


@app.get("/api/jobs/{job_id}/audio")
async def job_audio(job_id: str) -> FileResponse:
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.source != "audio" or not job.audio_path.exists():
        raise HTTPException(status_code=404, detail="Audio indisponible")
    suffix = job.audio_path.suffix.lower()
    media = {
        ".wav": "audio/wav",
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".ogg": "audio/ogg",
        ".flac": "audio/flac",
    }.get(suffix, "application/octet-stream")
    return FileResponse(str(job.audio_path), media_type=media, filename=job.audio_path.name)


@app.get("/api/jobs/{job_id}/turns")
async def job_turns(job_id: str) -> JSONResponse:
    """Renvoie le transcript par TOURS DE PAROLE (un objet par turn, avec
    timestamps précis + texte + label speaker) + le mapping facultatif
    SPEAKER_XX → nom affichable (édité par l'utilisateur via PATCH).

    Le frontend utilise les `start`/`end` pour synchroniser avec l'audio et
    le mapping pour afficher les vrais noms à la place des SPEAKER_00, etc.
    """
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    turns_path = job.out_dir / "turns.json"
    if not turns_path.exists():
        # Cas job ancien (avant garde du transcript) ou source=transcript.
        return JSONResponse({"turns": [], "speakers": {}, "hasTurns": False})
    import json as _json
    try:
        turns = _json.loads(turns_path.read_text(encoding="utf-8"))
    except (OSError, ValueError) as exc:
        log.warning("Lecture turns.json échouée pour %s : %s", job_id, exc)
        raise HTTPException(status_code=500, detail="turns.json illisible")
    speakers_path = job.out_dir / "speakers.json"
    speakers_map: dict = {}
    if speakers_path.exists():
        try:
            speakers_map = _json.loads(speakers_path.read_text(encoding="utf-8"))
        except (OSError, ValueError):
            speakers_map = {}
    return JSONResponse({
        "turns": turns,
        "speakers": speakers_map,
        "hasTurns": True,
    })


class SpeakersPatch(BaseModel):
    """Mapping SPEAKER_XX → nom affichable. Une chaîne vide ou null retire
    l'entrée (= retour au label brut SPEAKER_XX)."""
    updates: dict[str, str | None]


@app.patch("/api/jobs/{job_id}/speakers")
async def job_speakers_patch(job_id: str, payload: SpeakersPatch) -> JSONResponse:
    """Met à jour le mapping speakers.json côté backend (persistant, partagé
    entre la vue transcript et l'éditeur s'il vient à l'utiliser). Le
    frontend, lui, applique le mapping au rendu — turns.json reste source
    de vérité (labels SPEAKER_XX inchangés)."""
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if (job.out_dir / "turns.json").exists() is False:
        raise HTTPException(status_code=400, detail="Pas de turns disponibles pour ce job")

    import json as _json
    speakers_path = job.out_dir / "speakers.json"
    current: dict[str, str] = {}
    if speakers_path.exists():
        try:
            current = _json.loads(speakers_path.read_text(encoding="utf-8"))
        except (OSError, ValueError):
            current = {}

    for label, name in payload.updates.items():
        if name is None or not str(name).strip():
            current.pop(label, None)
        else:
            current[label] = str(name).strip()

    speakers_path.write_text(_json.dumps(current, ensure_ascii=False, indent=2),
                              encoding="utf-8")
    return JSONResponse({"speakers": current})


@app.get("/api/jobs/{job_id}/download")
async def download(job_id: str, kind: str = "report"):
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")

    if kind == "transcript":
        path = job.transcript_path
        if path is None or not path.exists():
            raise HTTPException(status_code=404, detail="Transcript indisponible")
        media = "text/plain; charset=utf-8"
    else:
        path = job.report_docx_path or job.report_path
        if path is None or not path.exists():
            raise HTTPException(status_code=404, detail="Fichier indisponible")
        media = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if path.suffix == ".docx"
            else "text/markdown"
        )
    filename = f"{job.label}_{path.name}" if job.label else path.name

    # Lit le fichier en mémoire (permet de le régénérer / relâcher le handle rapidement,
    # et surtout de renvoyer une erreur claire si Word / OneDrive le verrouille).
    try:
        data = path.read_bytes()
    except PermissionError as exc:
        log.warning("Téléchargement bloqué, fichier verrouillé : %s", exc)
        raise HTTPException(
            status_code=423,
            detail=(
                "Le fichier .docx est verrouillé par une autre application "
                "(Word ouvert ou synchro OneDrive). Fermez-le puis réessayez."
            ),
        ) from exc
    except OSError as exc:
        log.exception("Lecture du rapport impossible : %s", exc)
        raise HTTPException(status_code=500, detail="Lecture du rapport impossible") from exc

    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(content=data, media_type=media, headers=headers)


@app.patch("/api/jobs/{job_id}")
async def rename_job(job_id: str, body: RenamePayload) -> JSONResponse:
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.status in ("queued", "running"):
        raise HTTPException(status_code=409, detail="Traitement en cours")

    new_label = _sanitize_label(body.label)
    if not new_label:
        raise HTTPException(status_code=400, detail="Nom invalide")
    if new_label == job.label:
        return JSONResponse(job.public())

    new_folder = HISTORY_DIR / new_label
    if new_folder.exists():
        raise HTTPException(status_code=409, detail="Un dossier avec ce nom existe déjà")

    old_folder = job.out_dir
    try:
        old_folder.rename(new_folder)
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"Renommage impossible : {e}") from e

    job.out_dir = new_folder
    job.label = new_label
    job.audio_path = new_folder / job.audio_path.name
    if job.report_path:
        job.report_path = new_folder / job.report_path.name
    if job.report_docx_path:
        job.report_docx_path = new_folder / job.report_docx_path.name
    if job.transcript_path:
        job.transcript_path = new_folder / job.transcript_path.name
    return JSONResponse(job.public())


@app.patch("/api/jobs/{job_id}/report")
async def update_report(job_id: str, body: ReportPayload) -> JSONResponse:
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.status != "done":
        raise HTTPException(status_code=409, detail="Rapport non disponible")
    md_path = job.out_dir / "compte_rendu.md"
    md_path.write_text(body.markdown, encoding="utf-8")
    job.report_path = md_path
    job.report_markdown = body.markdown
    docx_path = job.out_dir / "compte_rendu.docx"
    try:
        _md_to_docx(md_path, docx_path)
        job.report_docx_path = docx_path
    except Exception as exc:
        log.warning("Regen DOCX échouée pour job %s : %s", job.id, exc)
    return JSONResponse(job.public())


@app.post("/api/jobs/{job_id}/open-folder")
async def open_job_folder(job_id: str) -> dict:
    """Ouvre le dossier de la réunion dans l'explorateur de fichiers.
    Le backend tourne en local sur la machine de l'utilisateur, il peut donc
    demander à l'OS d'ouvrir le dossier directement."""
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    folder = job.out_dir
    if not folder.exists():
        raise HTTPException(status_code=404, detail="Dossier introuvable")
    try:
        import subprocess
        if sys.platform == "win32":
            # explorer.exe (et non os.startfile) : un processus en
            # arrière-plan ne peut pas voler le premier plan, mais explorer
            # a ce droit et réutilise une fenêtre déjà ouverte sur ce
            # dossier — la fenêtre passe donc bien devant.
            subprocess.Popen(["explorer", str(folder)])  # noqa: S607
        elif sys.platform == "darwin":
            subprocess.Popen(["open", str(folder)])
        else:
            subprocess.Popen(["xdg-open", str(folder)])
    except Exception as exc:
        log.exception("Ouverture du dossier impossible : %s", exc)
        raise HTTPException(
            status_code=500, detail="Ouverture du dossier impossible"
        ) from exc
    return {"ok": True}


@app.delete("/api/jobs/{job_id}")
async def delete_job(job_id: str) -> dict:
    job = jobs.pop(job_id, None)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.status in ("queued", "running"):
        jobs[job_id] = job
        raise HTTPException(status_code=409, detail="Traitement en cours, impossible de supprimer")

    failures = _force_rmtree(job.out_dir)
    if job.out_dir.exists():
        # Échec : on remet le job en mémoire pour garder l'UI cohérente.
        jobs[job_id] = job
        log.warning("Suppression du dossier %s partielle/échouée : %s",
                    job.out_dir, failures)
        hint = failures[0] if failures else "dossier encore présent"
        raise HTTPException(
            status_code=423,
            detail=(
                f"Suppression impossible : {hint}. "
                "Fermez Word / l'Explorateur (ou attendez la fin de la synchro OneDrive) puis réessayez."
            ),
        )
    return {"ok": True}


# ── Dossiers (catégories = sous-dossiers réels de Documents/Réunions/) ──────
@app.get("/api/folders")
async def get_folders() -> dict:
    return {"folders": list_category_folders()}


@app.post("/api/folders")
async def create_folder(body: FolderPayload) -> dict:
    name = _sanitize_label(body.name)
    if not name:
        raise HTTPException(status_code=400, detail="Nom de dossier invalide")
    target = HISTORY_DIR / name
    try:
        target.mkdir(parents=True, exist_ok=False)
    except FileExistsError as exc:
        raise HTTPException(status_code=409, detail="Ce dossier existe déjà") from exc
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"Création impossible : {exc}") from exc
    return {"folders": list_category_folders()}


@app.delete("/api/folders/{name}")
async def delete_folder(name: str) -> dict:
    d = HISTORY_DIR / name
    if not d.is_dir() or _looks_like_meeting(d):
        raise HTTPException(status_code=404, detail="Dossier introuvable")
    has_meetings = any(
        s.is_dir() and _looks_like_meeting(s) for s in d.iterdir()
    )
    if has_meetings:
        raise HTTPException(
            status_code=409,
            detail="Le dossier contient des réunions — déplacez-les d'abord.",
        )
    _force_rmtree(d)
    if d.exists():
        raise HTTPException(
            status_code=423,
            detail="Suppression impossible (dossier verrouillé). Fermez l'Explorateur puis réessayez.",
        )
    return {"folders": list_category_folders()}


@app.post("/api/jobs/{job_id}/folder")
async def move_job_folder(job_id: str, body: MoveFolderPayload) -> JSONResponse:
    """Déplace le dossier de la réunion dans une catégorie (sous-dossier réel
    de Documents/Réunions/) ou la ramène à la racine (folder = null)."""
    job = jobs.get(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job introuvable")
    if job.status in ("queued", "running"):
        raise HTTPException(status_code=409, detail="Traitement en cours")

    raw = (body.folder or "").strip()
    cat = _sanitize_label(raw) if raw else None
    if raw and not cat:
        raise HTTPException(status_code=400, detail="Nom de dossier invalide")
    if cat == (job.folder or None):
        return JSONResponse(job.public())

    parent = HISTORY_DIR / cat if cat else HISTORY_DIR
    if cat:
        try:
            parent.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            raise HTTPException(status_code=500, detail=f"Dossier inaccessible : {exc}") from exc

    dest = parent / job.out_dir.name
    if dest.exists():
        raise HTTPException(
            status_code=409,
            detail="Une réunion du même nom existe déjà dans ce dossier",
        )
    try:
        shutil.move(str(job.out_dir), str(dest))
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"Déplacement impossible : {exc}") from exc

    job.out_dir = dest
    job.audio_path = dest / job.audio_path.name
    if job.report_path:
        job.report_path = dest / job.report_path.name
    if job.report_docx_path:
        job.report_docx_path = dest / job.report_docx_path.name
    if job.transcript_path:
        job.transcript_path = dest / job.transcript_path.name
    job.folder = cat
    return JSONResponse(job.public())
